#!/usr/bin/env python3
"""Build the Sensors-oriented (CAS Q3) review as Markdown and Word."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from refs import BIB

ROOT = Path(__file__).resolve().parent
FIG = ROOT / "figures"

_order: list[str] = []
_num: dict[str, int] = {}


def cite(*keys: str) -> str:
    nums = []
    for key in keys:
        if key not in BIB:
            raise KeyError(f"unknown citation key: {key}")
        if key not in _num:
            _order.append(key)
            _num[key] = len(_order)
        nums.append(_num[key])
    nums = sorted(set(nums))
    return "[" + ",".join(str(n) for n in nums) + "]"


def reset_cites() -> None:
    _order.clear()
    _num.clear()


def set_run_font(run, name="Times New Roman", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_text(p, text, **kwargs):
    run = p.add_run(text)
    set_run_font(run, **kwargs)
    return run


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        run.font.name = "Times New Roman"


def para(doc, text, *, first_line=True, italic=False, size=11, bold=False, center=False, space_after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if first_line and not center:
        p.paragraph_format.first_line_indent = Inches(0.3)
    add_text(p, text, size=size, italic=italic, bold=bold)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Inches(0)
    add_text(p, text, size=10, italic=True)
    return p


def add_table(doc, headers, rows, title):
    caption(doc, title)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_text(p, h, size=8, bold=True)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            add_text(p, val, size=8)
    doc.add_paragraph()


def add_figure(doc, path: Path, cap: str, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption(doc, cap)


def write_markdown(path: Path, blocks: list[tuple]) -> None:
    lines = []
    for kind, payload in blocks:
        if kind == "title":
            lines += [f"# {payload}", ""]
        elif kind == "authors":
            lines += [payload, ""]
        elif kind == "affil":
            lines += [payload, ""]
        elif kind == "corr":
            lines += [f"**{payload}**", ""]
        elif kind == "h1":
            lines += [f"## {payload}", ""]
        elif kind == "h2":
            lines += [f"### {payload}", ""]
        elif kind == "p":
            lines += [payload, ""]
        elif kind == "caption":
            lines += [f"*{payload}*", ""]
        elif kind == "fig":
            rel = payload.relative_to(path.parent)
            lines += [f"![]({rel.as_posix()})", ""]
        elif kind == "table":
            title, headers, rows = payload
            lines += [f"*{title}*", ""]
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in rows:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")
        elif kind == "refs":
            lines += ["## References", ""]
            for i, key in enumerate(_order, 1):
                lines.append(f"[{i}] {BIB[key]}")
                lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def manuscript_blocks() -> list[tuple]:
    reset_cites()
    b: list[tuple] = []

    b.append(("title", "Depth Vision Sensors: Technological Evolution, Enabling Algorithms, and Applications in Mechatronic Systems"))
    b.append(("authors", "Bo Zhang1,2 | Hongchao Cui1,2"))
    b.append(("affil", "1 School of Mechanical, Electronic and Control Engineering, Beijing Jiaotong University, Beijing, China\n2 Beijing Key Laboratory of Flow and Heat Transfer of Phase Changing in Micro and Small Scale, Beijing Jiaotong University, Beijing, China"))
    b.append(("corr", "Corresponding author: Hongchao Cui (hccui@bjtu.edu.cn)"))

    b.append(("h1", "Abstract"))
    b.append((
        "p",
        "Background: Active depth cameras and solid-state LiDAR now sit on industrial robots, indoor mobile bases, and small unmanned aerial vehicles (UAVs), yet designers still treat “depth sensor” as one specification. Methods: This structured narrative review screens English-language work from 2001 to early 2025 in IEEE Xplore, Scopus, and Web of Science on structured light, time-of-flight (ToF), RGB-D, solid-state LiDAR, depth completion, visual-inertial and RGB-D SLAM, bin picking, collaborative robots, and UAV ranging. Priority is given to hardware papers, comparative evaluations with reported error, and application studies that name the modality. Results: Structured light still owns close-range accuracy; compact iToF and active stereo own room-scale mobile heads; solid-state LiDAR owns outdoor metres. Those envelopes, not marketing generations, decide whether KinectFusion-style mapping, visual-inertial-depth odometry, LiDAR-inertial odometry, or a safety-rated stop is even feasible. Conclusions: A mechatronic stack should be chosen by range, lighting, surface, mass, and safety integrity, keep a metric low-level range channel when a motor can stop, and treat calibration as a first-class design item. The paper maps sensor physics to estimators and applications; it does not replace modality-specific optics surveys.",
    ))
    b.append((
        "p",
        "Keywords: depth vision sensors; mechatronic systems; structured light; time-of-flight; solid-state LiDAR; RGB-D; SLAM; sensor fusion; collaborative robots; UAV inspection",
    ))
    b.append(("h1", "Highlights"))
    b.append(("p", "• Hardware generations did not replace one another. They split the operating envelope: structured light for close-range accuracy, ToF and active stereo for compact room-scale sensing, solid-state LiDAR for outdoor metres."))
    b.append(("p", "• Comparative evaluations (Kinect v1/v2, Azure Kinect, RealSense D400-class, multi-camera indoor tests) give order-of-magnitude errors that are usable for system design if they are not treated as datasheet guarantees."))
    b.append(("p", "• The estimator must match the front-end: TSDF fusion for tabletops, visual-inertial-depth odometry for indoor bases, LiDAR-inertial odometry for sparse long-range clouds."))
    b.append(("p", "• A consumer RGB-D camera is not a safety-rated device. Collaborative cells still need ISO/TS 15066 logic, a safety controller, and usually a redundant channel."))
    b.append(("p", "• Event cameras, on-sensor AI, and photonic beam steering are real research directions. They are not yet drop-in replacements for frame-based RGB-D in a warehouse bill of materials."))
    b.append(("fig", FIG / "graphical_abstract.png"))
    b.append(("caption", "Graphical abstract. Active depth families feed distinct estimators and only then become mechatronic functions. Limits (sunlight, watts, calibration, fusion integrity) remain first-order."))

    b.append(("h1", "1. Introduction"))
    b.append((
        "p",
        f"Autonomous mobile robots, flexible manufacturing cells, and small UAVs all require a timely geometric model of the workspace {cite('cadena2016','floreano2015','villani2018','halmet2019')}. Among perception options, vision remains attractive because it is non-contact, relatively inexpensive, and information-dense. Two-dimensional cameras supply appearance. Depth vision adds metric range, which simplifies obstacle clearance, grasp planning, and map scale. In a mechatronic system that coupling is not cosmetic: the same millimetre that is harmless in a visualization can be the difference between a successful insertion and a jammed assembly, or between a UAV standoff and a girder strike.",
    ))
    b.append((
        "p",
        f"This review is restricted to active depth vision: sensors that illuminate the scene with a controlled optical source and recover three-dimensional geometry from the returned light. The main families are structured light and active stereo, continuous-wave and pulsed ToF cameras, and solid-state LiDAR {cite('zhang2012sl','geng2011','foix2011','horaud2016','ho2022','roriz2022')}. We contrast them with passive stereo and learning-based monocular depth, which rely on ambient texture or data-driven priors {cite('scharstein2002','eigen2014','godard2017')}. Passive methods have improved rapidly, but active sensors remain the default in many deployed mechatronic systems because they provide dense or semi-dense metric depth on weakly textured surfaces and under changing indoor lighting {cite('khoshelham2012','halmet2019','sarbolandi2015')}.",
    ))
    b.append((
        "p",
        f"The 2010 Microsoft Kinect showed that dense indoor depth could be obtained at consumer cost {cite('khoshelham2012','shotton2011')}. That device, and the algorithms built on it, moved RGB-D mapping and human pose estimation out of specialized laboratories {cite('newcombe2011','izadi2011','henry2012')}. Subsequent products reduced size and power (mobile ToF and active-stereo modules) and then extended range (solid-state LiDAR), which changed the set of machines that could carry a depth sensor {cite('keselman2017','gyongy2022','ho2022','roriz2022')}.",
    ))
    b.append((
        "p",
        f"Several surveys already cover pieces of this landscape (Table 1). Lock-in ToF cameras {cite('foix2011')}, ToF range imaging more broadly {cite('horaud2016','kolb2010','zanuttigh2016')}, structured-light metrology {cite('zhang2012sl','geng2011')}, comparative RGB-D tests {cite('halmet2019','giancola2018')}, visual SLAM {cite('cadena2016','taketomi2017','barros2022','wang2024slam')}, event cameras {cite('gallego2022')}, and LiDAR hardware {cite('ho2022','roriz2022','li2020lidar','raj2020')} are all better treated in those sources than they can be in one paper. What is still useful, and what this paper attempts, is not another first-principles optics tutorial. It is a system-level map that answers a design question: which physical generation of depth sensor made which mechatronic function practical, and which residual errors still dominate in the field.",
    ))
    b.append((
        "table",
        (
            "Table 1. Closest prior surveys and the slice added here. This paper is complementary, not a replacement.",
            ["Prior survey", "Main focus", "What a mechatronic designer still has to assemble"],
            [
                [f"Foix et al.; Horaud et al. {cite('foix2011','horaud2016')}", "ToF camera physics and calibration", "How iToF/dToF sit next to SL and LiDAR on a robot"],
                [f"Zhang; Geng {cite('zhang2012sl','geng2011')}", "Structured-light metrology", "When a fringe scanner is the wrong robot sensor"],
                [f"Halmetschlager-Funek et al. {cite('halmet2019')}", "Empirical RGB-D ranking", "Estimator and application consequences"],
                [f"Cadena et al.; later V-SLAM surveys {cite('cadena2016','barros2022')}", "SLAM algorithms", "Which depth front-end the estimator assumes"],
                [f"Ho et al.; Roriz et al. {cite('ho2022','roriz2022')}", "Solid-state / automotive LiDAR", "Indoor RGB-D and factory cells"],
                [f"Gallego et al. {cite('gallego2022')}", "Event cameras", "When event depth is worth the stack change"],
            ],
        ),
    ))

    b.append(("h2", "1.1 Scope and review method"))
    b.append((
        "p",
        "The paper is a structured narrative review. It is not a PRISMA systematic review and does not claim an exhaustive count of every RGB-D paper since 2010. Inventing a flow-diagram numerator would be dishonest. What we did specify is the search frame, the inclusion rule, and the reason a paper was kept.",
    ))
    b.append((
        "p",
        f"Sources. English-language records from approximately January 2001 to March 2025 were screened in IEEE Xplore, Scopus, and Web of Science, with publisher sites (IEEE, Springer, Elsevier, MDPI, Wiley, ACM) used to retrieve full texts. Core books and standards were added by citation chasing {cite('thrun2005','hansard2013','giancola2018','zanuttigh2016','iso15066','iso10218')}. Query stems included structured light, fringe projection, time-of-flight camera, RGB-D, Kinect, RealSense, Azure Kinect, solid-state LiDAR, optical phased array, depth completion, visual-inertial odometry, RGB-D SLAM, LiDAR-inertial odometry, bin picking, 6-DoF pose, collaborative robot, speed and separation monitoring, and UAV obstacle avoidance OR inspection. Boolean combinations were adapted to each database.",
    ))
    b.append((
        "p",
        "Inclusion. A record was kept if it (i) states a depth or ranging principle that a mechatronic designer can act on, (ii) reports a quantitative error, range, or system result, or (iii) is a widely cited survey that we needed in order to point readers to a deeper treatment. Exclusion. Purely biomedical or cinematic uses; papers that never name the depth modality; product blogs and unreviewed datasheets as primary evidence; and mechanical spinning LiDAR except as a baseline for solid-state designs. When two papers report the same device, the peer-reviewed metrology study is preferred over a manufacturer white paper. Table 1a records that rule.",
    ))
    b.append((
        "table",
        (
            "Table 1a. Inclusion and exclusion used in this structured narrative review.",
            ["Keep if", "Drop if", "How the record is used"],
            [
                ["Hardware principle or calibration is stated", "Modality never named", "Section 2 envelopes"],
                ["Peer-reviewed range or error numbers", "Datasheet-only claim with no method", "Tables 2-3"],
                ["Estimator assumes a stated depth front-end", "Algorithm paper with a generic “depth input”", "Section 3"],
                ["Application names the sensor family", "Market survey without a sensing requirement", "Section 4"],
                ["Standard that governs safety or robots", "Opinion piece", "Sections 4.4 and 5"],
            ],
        ),
    ))
    b.append((
        "p",
        f"This protocol is enough for a Sensors-style review to be reproducible in spirit. It is not enough to support a meta-analysis of depth RMSE across all consumer cameras, because experimental setups are not commensurate {cite('halmet2019','giancola2018','tolgyessy2021','kurillo2022')}. Non-English records and manufacturer white papers were not treated as primary evidence. The search was last refreshed in March 2025; later product SKUs are therefore outside the evidence window.",
    ))

    b.append(("h2", "1.2 Contributions and organization"))
    b.append((
        "p",
        f"The paper makes six contributions aimed at mechatronic designers rather than at optics specialists. First, it compares structured light, active stereo, iToF, dToF, and solid-state LiDAR by principle, reported range and error, and role. Second, it compiles device-level numbers from peer-reviewed evaluations of Kinect-class, RealSense-class, and Azure Kinect cameras so that Table 3 can be opened during a design review {cite('khoshelham2012','sarbolandi2015','carfagni2019','tolgyessy2021','kurillo2022')}. Third, it links raw artifacts to the estimators that machines actually run, including LiDAR-inertial and lidar-visual-inertial odometry for sparse long-range clouds {cite('zhang2014loam','shan2020','xu2022fastlio','shan2021lvisam')}. Fourth, it organizes applications by the depth property that is spent. Fifth, it states what standard benchmarks measure and what they do not (Table 8). Sixth, it gives an explicit selection flowchart and a checklist that can be copied into a requirements document. Section 2 reviews hardware. Section 3 reviews enabling algorithms. Section 4 reviews applications and selection. Section 5 discusses open problems. Section 6 concludes.",
    ))

    b.append(("h1", "2. Technological evolution of depth vision sensors"))
    b.append((
        "p",
        "A useful first cut is the operating envelope: the combination of range, spatial resolution, ambient-light tolerance, size, and cost that a machine can actually use. Figure 1 summarizes the three families that dominate current mechatronic designs.",
    ))
    b.append(("fig", FIG / "fig1_operating_envelopes.png"))
    b.append((
        "caption",
        "Figure 1. Typical operating envelopes of structured light / active stereo, ToF cameras, and solid-state LiDAR, together with the mechatronic roles that follow from those envelopes. Bounds are qualitative summaries of the literature in Table 2, not guarantees for a particular device.",
    ))

    b.append(("h2", "2.1 Structured light and active stereo"))
    b.append((
        "p",
        f"Structured-light sensors project a known pattern (dots, stripes, or time-varying fringes) and recover depth by triangulation between the projector and one or more cameras {cite('zhang2012sl','geng2011','salvi2010')}. If f is the focal length, B the baseline, and d the observed disparity, the pinhole relation is Z = fB/d. Error in disparity therefore grows into a range error that increases approximately with Z squared {cite('khoshelham2012')}. Pattern choice is itself a design variable: Salvi et al. classify time-multiplexed, neighbourhood, and direct-coding families, and show that no single code is optimal for both static metrology and a moving robot {cite('salvi2010')}. Industrial systems usually prefer sequential Gray-code plus phase-shifting fringes because the absolute unwrapping is robust and the phase gives sub-pixel disparity {cite('zhang2012sl','geng2011','zhang2018phase')}. Absolute-phase reviews make the same point in more optical detail: the unwrapping step, not the projector brightness, is usually what fails first on discontinuities {cite('zhang2018phase')}. Fourier-transform profilometry trades some robustness for a single-shot capture, which matters on moving parts {cite('su2001ftp')}. Industrial fringe-projection systems exploit this geometry at short range and can reach tens of micrometres to sub-millimetre accuracy on cooperative surfaces {cite('zhang2012sl')}. The price of that accuracy is a controlled standoff, a cooperative (or at least non-specular) surface, and a workspace that fits the calibrated volume.",
    ))
    b.append((
        "p",
        f"Consumer devices traded metrology-grade projectors for a static pseudo-random speckle and on-chip matching. Kinect v1 is the canonical example: a near-infrared (NIR) projector, an infrared camera, and a colour camera. Khoshelham and Elberink showed that its random depth error grows from a few millimetres near the sensor to about 4 cm near 5 m, while axial point spacing can reach about 7 cm at that range {cite('khoshelham2012')}. Early geometric studies of the same device reached the same qualitative conclusion and supplied calibration recipes that robotics groups still reuse {cite('smisek2011','herrera2012')}. Those numbers explain both the success of KinectFusion-style indoor reconstruction {cite('newcombe2011','izadi2011')} and the unsuitability of the same sensor as an outdoor navigation camera.",
    ))
    b.append((
        "p",
        f"Active stereo keeps triangulation but replaces a coded single-camera pattern with a stereo pair plus a texture projector. Intel RealSense R200/D400-class cameras follow this route {cite('keselman2017')}. Because matching runs on two infrared images, the system can still return depth when the projector is weak or switched off, which improves outdoor behaviour relative to first-generation speckle sensors {cite('keselman2017','halmet2019')}. The cost is compute for stereo matching and a stronger dependence on calibration {cite('scharstein2002','keselman2017')}. Indoor tests across structured-light, active-stereo, and ToF units confirm that no single consumer camera dominates bias, precision, lateral noise, lighting, and multi-sensor interference at once {cite('halmet2019','giancola2018')}.",
    ))

    b.append(("h2", "2.2 Time-of-flight cameras and mobile miniaturization"))
    b.append((
        "p",
        f"A ToF pixel estimates range from the travel time of light. Indirect ToF (iToF; lock-in or continuous-wave) measures the phase shift of a modulated source {cite('lange2001','foix2011','hansard2013')}. For a modulation frequency f_m the unambiguous range is on the order of c/(2 f_m); a typical 30 MHz tone wraps near 5 m, which is why multi-frequency operation appears on industrial iToF cameras {cite('foix2011','horaud2016','hansard2013')}. Direct ToF (dToF) timestamps a short pulse, often with single-photon avalanche diode (SPAD) arrays and time-to-digital converters {cite('niclass2013','gyongy2022','horaud2016')}. iToF became practical in compact cameras because a single illuminator and a single sensor replace a precision stereo baseline {cite('foix2011','sarbolandi2015')}. Reported indoor working distances for consumer iToF, including Kinect v2, are typically 0.5-4.5 m {cite('sarbolandi2015','lachat2015','pagliari2015')}. Side-by-side studies of Kinect v1 and v2 show that the ToF unit reduced some structured-light artifacts (texture dependence, multi-device interference) while introducing others (multipath in corners, flying pixels, wiggling) {cite('sarbolandi2015','lachat2015','pagliari2015','wasenmuller2017')}. Azure Kinect, the later Microsoft ToF camera, reduced random depth noise relative to Kinect v2 in near-FOV modes under about 3 m and improved spatial accuracy beyond about 2.5 m in laser-scanner comparisons, without removing sunlight or multipath as design limits {cite('tolgyessy2021','kurillo2022')}. Sunlight, multipath, and flying pixels remain first-order limitations {cite('sarbolandi2015','lachat2015','kadambi2013')}.",
    ))
    b.append((
        "p",
        f"The same physics explains the mobile-ToF wave of the 2010s. A phase camera does not need a wide mechanical baseline, so it fits a phone or an embedded robot head. Early mobile iToF modules were lower in spatial resolution than structured light but more convenient at 2-5 m for augmented reality and coarse scene layout {cite('horaud2016','gyongy2022')}. dToF SPAD arrays later improved ambient-light rejection and power per unit of range, at the expense of histogram memory and, often, spatial resolution {cite('gyongy2022')}. Gyongy, Dutton, and Henderson review the dToF signal chain and show why on-chip histogram compression, not only detector quantum efficiency, now limits array size {cite('gyongy2022')}. For a mechatronic designer the takeaway is narrower than the semiconductor literature: if the workspace has corners, glass, or two ToF units that can see each other, budget a multipath and interference test before you freeze the SKU {cite('kadambi2013','sarbolandi2015','halmet2019')}. Coded ToF can recover a time profile rather than a single phase, but the extra illumination and compute rarely fit a battery-powered head {cite('kadambi2013')}.",
    ))

    b.append(("h2", "2.3 Solid-state LiDAR and the long-range shift"))
    b.append((
        "p",
        f"Once the task leaves the room (warehouse aisles, yards, roads, bridge girders) consumer RGB-D cameras run out of photons and out of unambiguous range. Mechanical spinning LiDAR already solved long-range ranging. Solid-state and semi-solid designs try to keep that range while removing a bulky rotator {cite('ho2022','roriz2022','li2020lidar','raj2020','royo2019','behroozpour2017')}. MEMS mirrors scan a laser with a small moving mass and are the most commercially mature semi-solid option. Flash LiDAR illuminates a patch at once and is closer to a ToF camera, with range and resolution set by peak power and pixel count. Optical phased arrays and metasurface beam steerers aim at a fully solid-state scanner {cite('kim2021nano','park2021slm','sun2013opa','poulton2017','zhang2022mems','hsu2021opa')}. Large-scale silicon photonic arrays and MEMS-on-photonics demonstrators show that chip-scale beam steering is no longer only a laboratory sketch {cite('sun2013opa','zhang2022mems')}. Circuit- and architecture-level reviews stress that the scanner is only one block: laser, detector, timing, and eye-safety set the range as much as the steering method {cite('behroozpour2017','royo2019')}. For mechatronic integration the practical facts are simpler: solid-state units are smaller and potentially more robust to vibration, but they are still sparse compared with RGB-D, still expensive at automotive grade, and still require careful time synchronization with cameras and IMUs {cite('roriz2022','li2020lidar','debeunne2020')}. Automotive surveys also stress eye-safety, rain/fog backscatter, and the need for a perception stack that does not treat a sparse cloud as if it were a Kinect frame {cite('roriz2022','li2020lidar')}. A robot or UAV buyer should therefore pick the scanner class by payload and field of view first: MEMS for a compact outdoor head, flash for a short-range patch with no moving parts, and OPA or metasurface units only when the programme can absorb a research-grade integration {cite('ho2022','hsu2021opa','raj2020')}.",
    ))
    b.append((
        "p",
        "Table 2 collects order-of-magnitude figures reported in the literature. Values are typical envelopes, not guarantees for a particular serial number. Two design rules follow. First, do not treat \"depth camera\" as one specification: a sensor that is excellent for bin picking is usually the wrong sensor for a 30 m warehouse aisle. Second, hardware generations did not replace one another; they split the operating envelope. Structured light still wins close-range accuracy. ToF wins compactness at room scale. Solid-state LiDAR wins range. Mechatronic architectures increasingly carry more than one of them.",
    ))
    b.append((
        "table",
        (
            "Table 2. Representative operating envelopes of active depth sensors used in mechatronics. Numbers are typical ranges reported in peer-reviewed evaluations or hardware surveys; product variants differ.",
            ["Family", "Principle", "Typical usable range", "Reported depth error (order)", "Spatial sampling", "Main field weakness", "Typical role", "Sources"],
            [
                ["Structured light (consumer)", "Static-pattern triangulation", "0.5-4 m", "mm near field; ~4 cm random error near 5 m (Kinect v1)", "Dense VGA-class", "Sunlight; shiny or absorbing surfaces", "Indoor mapping, HRI prototypes", cite("khoshelham2012", "halmet2019", "sarbolandi2015")],
                ["Fringe / industrial SL", "Phase-shifting triangulation", "0.1-2 m (setup-dependent)", "tens of um to sub-mm", "Very dense", "Workspace size; ambient light", "In-line metrology, bin picking", cite("zhang2012sl", "geng2011")],
                ["Active stereo", "IR stereo + texture projector", "0.2-10 m (best ~0.3-3 m)", "cm-level, lighting-dependent", "Dense HD-class", "Calibration; compute; texture holes", "Indoor / semi-outdoor robots", cite("keselman2017", "halmet2019")],
                ["iToF camera", "CW phase", "0.5-5 m", "cm-level; multipath bias", "Dense QVGA-VGA", "Multipath; sunlight; flying pixels", "Mobile robots, mid-range HRI", cite("foix2011", "horaud2016", "sarbolandi2015", "lachat2015")],
                ["dToF SPAD", "Pulsed photon timing", "~1-10 m consumer; longer if laser-class", "cm-level; better ambient rejection than iToF", "Often coarser", "Histogram memory; pile-up", "AR, short-range robots", cite("gyongy2022", "niclass2013")],
                ["Solid-state / MEMS LiDAR", "Scanned or flash ToF", "10-200+ m", "few cm (range- and target-dependent)", "Sparse points", "Cost; weather; boresight", "Outdoor AMR, UAV, vehicle", cite("ho2022", "roriz2022", "li2020lidar")],
            ],
        ),
    ))
    b.append((
        "p",
        f"A practical selection sequence used in many robot shops is: write down the minimum and maximum range, the worst lighting, the worst surface, the allowed mass and watts, and the safety integrity level. Only then open a catalogue (Figure 2). Indoor rooms with matte walls still suit RGB-D {cite('khoshelham2012','keselman2017','halmet2019')}. Fixed cells with metal parts still suit industrial structured light {cite('zhang2012sl')}. Yards and roads still suit LiDAR, usually fused with cameras {cite('roriz2022','debeunne2020','yeong2021')}. If two of those worlds appear on one machine, budget two sensors and a calibration procedure, not a single \"universal\" depth camera.",
    ))
    b.append(("fig", FIG / "fig4_selection_flowchart.png"))
    b.append((
        "caption",
        "Figure 2. Selection from the operating envelope. The three exit boxes are families, not product SKUs. Mixed worlds require mixed sensors.",
    ))
    b.append((
        "p",
        f"Three short examples make the same point. A bin-picking cell with a 1.2 m standoff and oily steel parts should start from industrial structured light, a 6-DoF pose estimator, and a compliant grasp, not from a 100 m LiDAR {cite('zhang2012sl','correll2018','tenpas2017')}. An indoor delivery base that must see a pallet toe and a person in a corridor should start from a wide RGB-D camera plus a two-dimensional safety LiDAR and a visual-inertial estimator {cite('keselman2017','fankhauser2015','campos2021')}. A bridge-inspection UAV that must hold a 5-15 m standoff in wind should start from a lightweight solid-state ranger, an IMU, and a compact RGB inspector, and should not expect a phone-class ToF module to carry the ranging load {cite('floreano2015','spencer2019','ho2022')}.",
    ))

    b.append(("h2", "2.4 Device-level numbers and the function each generation unlocked"))
    b.append((
        "p",
        f"Family envelopes are not enough when a buyer has to choose a SKU. Table 3 lists numbers that peer-reviewed evaluations actually measured, not marketing ranges. Kinect v1 remains the best-documented structured-light consumer camera {cite('khoshelham2012','smisek2011')}. Kinect v2 and Azure Kinect document the ToF path, including the fact that Azure Kinect is not uniformly better than v2 at every distance and FOV mode {cite('sarbolandi2015','wasenmuller2017','tolgyessy2021','kurillo2022')}. RealSense SR300 and D415 document the coded-light and active-stereo Intel line with metrological characterizations rather than blog tests {cite('carfagni2017','carfagni2019','keselman2017')}. Halmetschlager-Funek et al. remain the widest indoor multi-device comparison and should be read before any single-camera ranking is trusted {cite('halmet2019')}.",
    ))
    b.append((
        "table",
        (
            "Table 3. Device-level figures taken from peer-reviewed evaluations. These are experimental orders of magnitude, not acceptance specifications.",
            ["Device (study)", "Principle", "Reported range or setup", "Reported error / noise", "Note for designers"],
            [
                [f"Kinect v1 {cite('khoshelham2012')}", "Structured light", "Indoor, to ~5 m", "Random error few mm near field; ~4 cm near 5 m; axial spacing ~7 cm at 5 m", "Good enough for indoor TSDF; not a survey instrument"],
                [f"Kinect v2 {cite('sarbolandi2015','wasenmuller2017','lachat2015')}", "iToF", "Typically 0.5-4.5 m indoor", "Better texture robustness than v1; multipath, flying pixels, wiggling", f"Indoor robot head; model the bias {cite('fankhauser2015')}"],
                [f"Azure Kinect {cite('tolgyessy2021','kurillo2022')}", "iToF", "Mode-dependent; studies cover ~0.5-5 m", "NFOV noise ~2x smaller than v2 under 3 m; better than v2 beyond ~2.5 m in one laser-scanner test", "Check FOV mode; not a sunlight camera"],
                [f"RealSense SR300 {cite('carfagni2017')}", "Coded light", "Short-range desktop", "Metrological characterization; close-range only", "Do not stretch to navigation"],
                [f"RealSense D415 {cite('carfagni2019')}", "Active stereo", "Desktop / robot cell distances in the study", "cm-level, setup-dependent", "Calibration and lighting dominate"],
                [f"RealSense D400 family {cite('keselman2017')}", "Active stereo", "Roughly 0.2-10 m; best nearer", "Matching artefacts if texture and projector are both weak", "Outdoor-capable relative to speckle SL"],
                [f"Ten-camera indoor test {cite('halmet2019')}", "SL / stereo / ToF mix", "Robot-cell indoor grid", "No single winner on bias, precision, lighting, multi-device", "Re-rank after your lighting and material"],
            ],
        ),
    ))
    b.append((
        "p",
        "Figure 3 states the causal claim of this review in one picture: each hardware step unlocked a function that was previously impractical at that size and cost. It did not make the previous function obsolete.",
    ))
    b.append(("fig", FIG / "fig5_hardware_functions.png"))
    b.append((
        "caption",
        "Figure 3. Hardware steps and the mechatronic functions they made newly practical. Emerging event and photonic devices are shown as research, not as current bill-of-materials items.",
    ))

    b.append(("h1", "3. Enabling technologies: from raw depth to a state estimate"))
    b.append((
        "p",
        "Raw depth is not a pose, a mesh, or a grasp. This section reviews the software layers that mechatronic systems insert between the sensor and the controller.",
    ))

    b.append(("h2", "3.1 Depth data processing and enhancement"))
    b.append((
        "p",
        f"Consumer and automotive depth images share a small set of artifacts: impulse noise, invalid pixels (holes) on occlusions, black or specular surfaces, flying pixels on depth edges, and ToF multipath {cite('foix2011','halmet2019','sarbolandi2015')}. Classical spatial median filters and temporal averaging remain the first firmware line of defence. When a registered RGB image is available, joint bilateral or guided filtering uses colour edges to protect object boundaries while smoothing planar interiors {cite('kopf2007')}. These filters are cheap enough for embedded GPUs. They do not invent missing geometry.",
    ))
    b.append((
        "p",
        f"Depth completion does. Early methods inpainted by diffusion or multi-view geometry. Learning-based completion trained on RGB-D or RGB-plus-sparse-LiDAR pairs now dominates the literature {cite('silberman2012','eigen2014','laina2016','godard2017','godard2019','uhrig2017','ma2018','zhang2018ddc','cheng2018cspn','park2020nlspn','qiu2019deeplidar')}. Eigen et al. showed that a multi-scale network can predict depth from a single RGB image {cite('eigen2014')}; later residual and self-supervised models reduced the need for dense ground truth {cite('laina2016','godard2017','godard2019')}. Indoor work still leans on NYU-Depth v2-style labelled apartments {cite('silberman2012')}, and later on ScanNet- and Matterport-scale reconstructions when a method needs more than a few dozen rooms {cite('dai2017scannet','chang2017matterport')}. Zhang and Funkhouser showed that even a single RGB-D frame with holes can be completed if surface normals and occlusion boundaries are used as intermediate cues {cite('zhang2018ddc')}. Outdoor work leans on KITTI-style projected LiDAR, which is sparse and biased toward the road plane {cite('geiger2012kitti','geiger2013kitti','uhrig2017')}. For that setting, Uhrig et al. introduced sparsity-invariant convolutions {cite('uhrig2017')}, and Ma and Karaman showed that even a few hundred metric samples plus RGB yield a dense metric map {cite('ma2018')}. Spatial-propagation networks (CSPN, NLSPN) and DeepLiDAR-style normal guidance then became the standard way to spread those sparse metric seeds {cite('cheng2018cspn','park2020nlspn','qiu2019deeplidar')}. Later non-local and transformer architectures improve large holes by using long-range context, at a cost that still challenges small onboard computers {cite('guo2021pc')}. A mechatronic reading of these papers is that completion is a good virtual sensor for visualization and mid-level planning, and a bad sole input to a safety-rated stop. KITTI ranks are especially easy to over-read: a method that wins on projected road-plane LiDAR can still invent a thin pole or a glass door that a warehouse AMR must not hit {cite('uhrig2017','geiger2013kitti')}.",
    ))
    b.append((
        "p",
        "Table 4 rates common artifacts by industrial maturity. The important systems point is that learning-based completion is not yet a drop-in safety sensor: it can look plausible while being metrically wrong on transparent, thin, or never-seen objects.",
    ))
    b.append((
        "table",
        (
            "Table 4. Common depth artifacts and the maturity of current mitigations.",
            ["Artifact", "Dominant cause", "Typical mitigation", "Maturity"],
            [
                ["Impulse noise", "Low SNR, scattering", "Median filter; temporal averaging", "Mature (on-chip / firmware)"],
                ["Holes / invalid pixels", "Occlusion, IR absorption, specularity, out-of-range", "RGB-guided completion; multi-view fusion", "Emerging for real-time embedded use"],
                ["Flying pixels / motion blur", "Finite integration time", "Shorter exposure; IMU-aided deblur; event sensing", "Mixed: IMU fusion common; event methods still research"],
                ["Multipath (ToF)", "Indirect returns in corners and shiny rooms", "Multi-frequency modulation; coded ToF; model-based subtraction", f"Maturing on high-end iToF; still hard on cheap single-frequency units {cite('foix2011','kadambi2013')}"],
                ["Sunlight saturation", "Ambient NIR swamping the source", "Narrow-band filters; higher peak power; dToF / LiDAR; radar fill-in", f"Partly solved at LiDAR-class power; unsolved for cheap RGB-D outdoors {cite('halmet2019','sarbolandi2015')}"],
            ],
        ),
    ))

    b.append(("h2", "3.2 Sensor fusion and state estimation"))
    b.append((
        "p",
        "A depth camera alone is a poor navigation sensor: it is biased, partially observed, and asynchronous with the actuators. Fusion with an inertial measurement unit (IMU), wheel odometry, and/or a colour camera is therefore the default architecture on mobile robots and UAVs (Figure 4).",
    ))
    b.append(("fig", FIG / "fig2_sensor_fusion_stack.png"))
    b.append((
        "caption",
        "Figure 4. A typical estimation stack on a mobile robot. Depth is one residual source among several. The optional dense or semantic layer is expensive and is often omitted on large-scale outdoor maps.",
    ))
    b.append((
        "p",
        f"Loose coupling treats each sensor as a black-box pose or twist and fuses them in an extended or unscented Kalman filter. The implementation is modular and cheap; the estimate is statistically suboptimal and cannot correct inner calibration errors {cite('thrun2005','lynen2013')}. Lynen et al. showed that a modular multi-sensor filter is enough to fly a MAV if each sensor is treated as a delayed pose update {cite('lynen2013')}. Tight coupling, as in keyframe visual-inertial odometry (VIO), jointly optimizes visual residuals and IMU preintegration in a sliding window {cite('leutenegger2015','forster2017','qin2018','bloesch2015','huang2019vin','geneva2020')}. The same idea has a longer visual-odometry lineage: MonoSLAM and PTAM showed that a single camera can track a sparse map in real time; SVO and DSO later made semi-direct and direct monocular odometry fast enough for small platforms {cite('davison2007','klein2007','forster2014svo','engel2018dso')}. OpenVINS and related platforms made the inertial version reproducible {cite('geneva2020','huang2019vin')}. Adding metric depth yields visual-inertial-depth odometry: depth removes the monocular scale-gauge ambiguity and supplies geometric constraints in low-texture rooms where VIO would otherwise drift or fail {cite('cadena2016','newcombe2011','qin2018')}.",
    ))
    b.append((
        "p",
        f"Dense RGB-D fusion, from KinectFusion through ElasticFusion and BundleFusion, maintains a truncated signed distance function (TSDF) or surfel map for close-range interaction {cite('newcombe2011','whelan2015','dai2017')}. The volumetric idea is older than Kinect: Curless and Levoy already fused range images into a signed-distance volume {cite('curless1996')}. The KinectFusion loop is still the mental model: track the live depth against the model, integrate the new frame into a volume, and raycast a synthetic depth for the next track {cite('newcombe2011','izadi2011')}. DTAM showed that a similar dense track can be run from a single moving RGB camera if a photometric cost is optimized on a GPU {cite('newcombe2011dtam')}. ElasticFusion dropped the explicit pose graph in favour of dense deformation; BundleFusion restored global consistency with on-the-fly reintegration {cite('whelan2015','dai2017')}. Those maps are excellent for a tabletop and expensive for a warehouse. Graph-based SLAM adds loop closures so that local odometry does not accumulate without bound {cite('cadena2016','thrun2005')}. RGB-D SLAM systems and the TUM RGB-D benchmark made this stack reproducible {cite('henry2012','endres2014','sturm2012','kerl2013')}. Kerl et al. showed that a dense photometric+depth residual already works on a CPU-era RGB-D camera if exposure is modelled {cite('kerl2013')}. Later systems (ORB-SLAM, ORB-SLAM2/3, Kimera) combine sparse or semi-dense tracking with optional metric-semantic mapping and multi-robot extensions {cite('murartal2015','murartal2017','campos2021','rosinol2020','tian2022')}. ORB-SLAM3 in particular folded visual-inertial and multi-map operation into a library that many robot teams actually run {cite('campos2021')}. Direct methods such as LSD-SLAM and DSO showed that photometric residuals can replace or complement sparse features when exposure is well modelled {cite('engel2014','engel2018dso')}. Learned dense trackers such as DROID-SLAM, and neural implicit maps (iMAP, NICE-SLAM, Point-SLAM), raise map fidelity again, at a compute cost that still sits above most embedded robot computers {cite('teed2021','sucar2021','zhu2022nice','sandstrom2023')}.",
    ))
    b.append((
        "p",
        f"Outdoors, the front-end is usually LiDAR, not RGB-D. LOAM and LeGO-LOAM established feature-based lidar odometry {cite('zhang2014loam','shan2018')}. Tight lidar-inertial systems (FAST-LIO, LIO-SAM, FAST-LIO2) then made solid-state and spinning clouds usable on UAVs and handheld platforms without a separate visual pipeline {cite('xu2021fastlio','shan2020','xu2022fastlio')}. FAST-LIO2 in particular registers raw points into an incremental k-d tree and has been demonstrated on solid-state LiDARs with a small field of view {cite('xu2022fastlio')}. When texture and structure are both available, lidar-visual-inertial estimators (V-LOAM, LVI-SAM, R3LIVE) colour the same map and survive brief LiDAR or camera dropouts better than a single-modality filter {cite('zhang2015vloam','shan2021lvisam','lin2022r3live','debeunne2020')}. Fusion surveys document the complementary pair: cameras for semantics and texture, LiDAR for metric structure and lighting invariance {cite('debeunne2020','yeong2021')}.",
    ))
    b.append((
        "p",
        f"Point-cloud infrastructure (PCL, iterative closest point, OctoMap, Voxblox) determines whether these estimators can run online on a real chassis {cite('rusu2011','besl1992','hornung2013','oleynikova2017')}. ICP itself has a large variant tree: point-to-plane and generalized-ICP improve robustness on structured rooms, and comparative tests show that the variant, the sampling, and the outlier rule matter as much as the textbook algorithm name {cite('rusinkiewicz2001','segal2009gicp','pomerleau2013')}. For control engineers the interface that matters is usually a pose with covariance plus a costmap, not a research SLAM paper. Table 5 compares the fusion patterns that those libraries usually sit under.",
    ))
    b.append((
        "table",
        (
            "Table 5. Fusion patterns used in mechatronic state estimation.",
            ["Pattern", "Typical inputs", "Estimator", "Output", "Fits", "Main cost"],
            [
                ["Loose coupling", "Depth pose, IMU, wheels", "EKF / UKF", "6-DoF pose", "Structured indoor bases", "Easy, less accurate"],
                ["Tight VIO / VIDO", "RGB, IMU, depth", "Sliding-window BA", "Pose + sparse map", "Drones, AR, agile bases", "Calibration and CPU"],
                ["Dense RGB-D", "Depth + RGB", "TSDF / surfels", "Mesh or volume", "Manipulation, inspection", "Memory and scale"],
                ["Pose-graph SLAM", "Any odometry + loops", "Graph optimization", "Globally consistent map", "Buildings, multi-floor", "Place recognition"],
                ["Visual-LiDAR", "Camera + LiDAR + IMU", "Tight or loosely fused SLAM", "Large-scale metric map", "Outdoor AMR / UAV", f"Sync and extrinsics {cite('debeunne2020','zhang2015vloam')}"],
                ["LiDAR-inertial", "LiDAR + IMU", "iEKF or factor graph", "Pose + sparse map", "UAV, handheld SSL", f"Needs a structured cloud {cite('shan2020','xu2021fastlio','xu2022fastlio')}"],
                ["Lidar-visual-inertial", "LiDAR + camera + IMU", "Tight factor graph", "Coloured metric map", "Outdoor inspection", f"Sync and exposure {cite('shan2021lvisam','lin2022r3live')}"],
            ],
        ),
    ))

    b.append(("h2", "3.3 From geometry to semantics"))
    b.append((
        "p",
        f"Once a metric cloud exists, the next questions are what the object is and where the robot may go. Classical 3D detectors used hand-crafted histograms. Current networks consume raw points, voxels, or range images {cite('qi2017pointnet','qi2017pn2','guo2021pc')}. Semantic segmentation labels every point; instance segmentation separates two chairs of the same class. RGB remains useful for texture and class priors; depth remains useful for scale and occlusion. Pose-estimation benchmarks such as BOP and methods such as PoseCNN made 6-DoF object pose a measurable industrial problem rather than a one-off demonstration {cite('hodan2018','xiang2018')}.",
    ))
    b.append((
        "p",
        f"Beyond labels, mechatronic systems care about affordances: surface normals and curvature for grasp stability, traversable floor versus negative obstacles, and dynamic scene graphs that persist over time {cite('rosinol2020','tian2022','tenpas2017')}. These layers are not unique to depth sensors, but they become cheaper when scale is observed rather than inferred. In dynamic scenes, geometric SLAM without semantics remains brittle; recent surveys document the shift toward semantic and robust dynamic SLAM {cite('wang2024slam','placed2023')}.",
    ))

    b.append(("h2", "3.4 Geometric calibration as a first-class subsystem"))
    b.append((
        "p",
        f"Every number in Tables 2 and 3 assumes that the depth camera, the colour camera, and any IMU or LiDAR share a known frame. Zhang's planar calibrator is still the colour-camera starting point {cite('zhang2000')}. RGB-D devices need a joint model of colour intrinsics, depth distortion, and the depth-to-colour extrinsics; Herrera, Kannala, and Heikkila gave the recipe that most robotics stacks still follow {cite('herrera2012')}. Kinect-era papers already showed that an uncalibrated factory model leaves a systematic bowl in the point cloud {cite('khoshelham2012','smisek2011')}. On a mobile robot the same error looks like a sloping floor and a biased obstacle range {cite('fankhauser2015')}. Visual-inertial and lidar-inertial estimators can absorb a slowly varying extrinsic, but they cannot invent a time stamp that the hardware never sent {cite('qin2018','geneva2020','xu2022fastlio')}. A design review that skips the calibration and synchronization budget is not finished.",
    ))

    b.append(("h2", "3.5 Benchmarks and what they do not measure"))
    b.append((
        "p",
        f"A Sensors review that only names algorithms without naming the datasets they were scored on is hard to use. Table 8 lists the public suites that most of the papers in Sections 3.1-3.3 actually report. TUM RGB-D and ICL-NUIM made indoor RGB-D odometry comparable {cite('sturm2012','handa2014icl')}. EuRoC and TUM VI did the same for visual-inertial estimators on MAVs and handheld sensors {cite('burri2016euroc','schubert2018tumvi')}. KITTI and later driving sets (including nuScenes) dominate outdoor completion and lidar-camera fusion {cite('geiger2012kitti','geiger2013kitti','uhrig2017','caesar2020')}. NYU-Depth, ScanNet, and Matterport3D supply indoor labels for completion and semantics {cite('silberman2012','dai2017scannet','chang2017matterport')}. YCB and BOP supply objects and 6-DoF pose protocols for grasping cells {cite('calli2015ycb','hodan2018')}.",
    ))
    b.append((
        "p",
        f"Two caveats matter more than the leaderboard. First, a good absolute trajectory error on TUM RGB-D or EuRoC does not certify a warehouse aisle or a sunny yard: those sequences are short, mostly well lit, and free of the safety integrity requirement {cite('sturm2012','burri2016euroc','halmet2019')}. Second, completion RMSE on KITTI is a score on projected LiDAR, not a guarantee that thin, transparent, or never-seen objects are recovered {cite('uhrig2017','geiger2013kitti')}. A design review should therefore treat Table 8 as a map of where a method was tested, not as a substitute for a factory or flight trial.",
    ))
    b.append((
        "table",
        (
            "Table 8. Public benchmarks used by the algorithms reviewed here, and the mechatronic question each suite cannot answer.",
            ["Suite", "Typical front-end", "What it scores", "What it does not score"],
            [
                [f"TUM RGB-D; ICL-NUIM {cite('sturm2012','handa2014icl')}", "RGB-D", "Indoor ATE / reconstruction", "Sunlight, rain, safety integrity"],
                [f"EuRoC; TUM VI {cite('burri2016euroc','schubert2018tumvi')}", "Mono/stereo + IMU", "MAV / handheld VIO drift", "Warehouse aisles; ToF multipath"],
                [f"KITTI; KITTI depth {cite('geiger2012kitti','geiger2013kitti','uhrig2017')}", "Camera + spinning LiDAR", "Driving odometry and completion", "Indoor cells; glass and thin poles as safety objects"],
                [f"nuScenes {cite('caesar2020')}", "Camera + LiDAR + radar", "Multimodal driving perception", "Factory lighting; cobot cells"],
                [f"NYU; ScanNet; Matterport3D {cite('silberman2012','dai2017scannet','chang2017matterport')}", "RGB-D", "Indoor depth and semantics", "Metric safety; outdoor range"],
                [f"YCB; BOP {cite('calli2015ycb','hodan2018')}", "RGB-D on tabletop objects", "6-DoF pose and grasp datasets", "Oily steel bins; line vibration"],
            ],
        ),
    ))

    b.append(("h1", "4. Applications in mechatronic systems"))
    b.append((
        "p",
        "The same sensor physics appears as different requirements once a machine has a job. This section organizes applications by the depth property that is actually spent.",
    ))

    b.append(("h2", "4.1 Industrial automation and precision manufacturing"))
    b.append((
        "p",
        f"Bin picking and unstructured part handling spend spatial resolution and short-range accuracy. A sensor above the bin produces a point cloud; a pose estimator returns a 6-DoF grasp; a compliant gripper absorbs the residual error {cite('tenpas2017','mahler2017','zeng2017','hodan2018','xiang2018','correll2018','morrison2018','fang2020graspnet','wang2019densefusion')}. Grasp-from-point-cloud methods (GPD, Dex-Net, GG-CNN, GraspNet) showed that a depth image plus analytic or learned grasp scores can propose suction or parallel-jaw contacts without a full CAD model of every SKU {cite('tenpas2017','mahler2017','morrison2018','fang2020graspnet')}. When the part identity is known, multi-view pose networks of the Amazon Picking Challenge generation, PoseCNN, DenseFusion, and later estimators scored on BOP turn the same cloud into a 6-DoF object pose {cite('zeng2017','xiang2018','wang2019densefusion','hodan2018')}. The YCB object set made those scores comparable across laboratories {cite('calli2015ycb')}. The first Amazon Picking Challenge made the systems nature of this problem obvious: teams that treated perception, planning, and gripping as separately optimized modules underperformed relative to tightly integrated stacks, and suction often beat anthropomorphic hands {cite('correll2018')}. Structured light and laser triangulation still dominate when parts are metallic, overlapping, and closer than about two metres {cite('zhang2012sl','correll2018')}. RGB is typically fused for class identity; force and torque close the last millimetres. A designer should therefore budget the sensor for the worst SKU in the bin, not the average YCB object: black rubber and oily steel still empty a consumer RGB-D cloud {cite('halmet2019','calli2015ycb')}.",
    ))
    b.append((
        "p",
        f"In-line metrology spends repeatability and speed. A fringe-projection or ToF snapshot is compared with a computer-aided design (CAD) model. Fringe systems own the micrometre-to-sub-millimetre band {cite('zhang2012sl','geng2011')}. ToF cameras are used when the assembly is large and a few millimetres of error are acceptable, because a full-field frame arrives without a scanning gantry {cite('foix2011','horaud2016')}. Vibration of the line is then a fusion problem (encoders and triggers), not only an optics problem.",
    ))
    b.append((
        "p",
        f"Automated guided vehicles (AGVs) and autonomous mobile robots (AMRs) spend field of view and vertical coverage. A two-dimensional safety LiDAR sees a plane; a wide RGB-D camera or solid-state LiDAR sees pallets, overhangs, and people who lean into the aisle {cite('halmet2019','yeong2021','fankhauser2015')}. Table 6 summarizes the pairing.",
    ))
    b.append((
        "table",
        (
            "Table 6. Industrial uses mapped to sensing requirements.",
            ["Task", "Preferred depth family", "First-order requirement", "Usual extra sensors"],
            [
                ["Bin picking", "Industrial SL / laser triangulation", "Sub-mm to few-mm, high density", "RGB; force/torque"],
                ["Assembly guidance", "SL or stereo at fixed standoff", "Low latency, stable extrinsics", "Joint encoders; IMU"],
                ["Full-field inspection", "Fringe (precise) or ToF (fast)", "Repeatability, line triggering", "Conveyor encoder; RGB"],
                ["AMR 3D obstacle sense", "Wide RGB-D and/or solid-state LiDAR", "Field of view, frame rate, human safety", "2D safety LiDAR; IMU; wheels"],
            ],
        ),
    ))

    b.append(("h2", "4.2 Autonomous mobile robots for service and logistics"))
    b.append((
        "p",
        f"Indoor service robots made RGB-D SLAM a product feature. Devices such as RealSense D400-class cameras are chosen because they balance size, power, software support, and indoor range {cite('keselman2017','halmet2019')}. Keselman et al. documented the optical and matching behaviour of the R200 and D400 families, including the fact that projector texture is an aid rather than a requirement {cite('keselman2017')}. The algorithmic backbone is well documented: RGB-D odometry, loop closure, and a metric map {cite('henry2012','endres2014','sturm2012','murartal2017')}. Fankhauser et al. modelled Kinect v2 specifically as a navigation sensor and showed where its systematic bias matters for terrain {cite('fankhauser2015')}. Occupancy and signed-distance maps (OctoMap, Voxblox) then turn that cloud into a costmap a local planner can query {cite('hornung2013','oleynikova2017')}. The remaining failures are physical. Stairs and negative obstacles (curbs, open shafts) are invisible to a waist-height two-dimensional LiDAR. A forward depth camera can reconstruct them if the near field is valid and the robot is slow enough for the integration time {cite('fankhauser2015')}. Outdoor last-metre robots meet sunlight and long range, which is why the same software stack often grows a LiDAR or radar {cite('debeunne2020','yeong2021')}. Teams that evaluate only on TUM RGB-D or EuRoC should therefore add a sunlight-and-glass trial before they call the stack warehouse-ready {cite('sturm2012','burri2016euroc','halmet2019')}.",
    ))
    b.append((
        "p",
        f"Human-aware navigation is a semantic layer on the same geometry. Depth supplies a metric body hull; RGB pose estimators supply intent cues {cite('shotton2011','sarbolandi2015')}. Socially acceptable clearance is then a planner parameter, not a new sensor. Active SLAM surveys discuss how a robot should move to keep that map healthy, which is a planning problem built on the same depth front-end {cite('placed2023')}.",
    ))

    b.append(("h2", "4.3 UAVs and aerial robotics"))
    b.append((
        "p",
        f"On a small UAV the payload budget is measured in tens of grams and a few watts {cite('floreano2015','kumar2012uav')}. That constraint, more than any algorithm, explains why lightweight depth cameras and solid-state LiDARs appear on inspection drones while automotive-grade spinning units do not. System papers on vision-controlled micro aerial vehicles already treated cameras, IMUs, and a metric map as one stack rather than as an optional perception add-on {cite('scaramuzza2014','kumar2012uav','lynen2013')}.",
    ))
    b.append((
        "p",
        f"Obstacle avoidance is the first use. Depth gives a metric stop distance in GPS-denied corridors, forests, or plant rooms {cite('floreano2015','muller2023','barry2018')}. Barry, Florence, and Tedrake demonstrated tree avoidance at up to 14 m/s with a pushbroom stereo front-end running at 120 Hz on a small airframe {cite('barry2018')}. That result is a reminder that geometric ranging, not a particular branded depth camera, is the requirement; stereo, RGB-D, and LiDAR are interchangeable only after latency, weight, and lighting are checked. Event-camera and event-stereo systems push latency into the microsecond regime for fast approach speeds {cite('gallego2022','falanga2020','zhou2021evo','he2024')}. Falanga et al. used event cameras to dodge dynamic obstacles that a frame-based pipeline would blur {cite('falanga2020')}. Miniature platforms have demonstrated onboard depth-based avoidance with tight compute envelopes {cite('muller2023')}.",
    ))
    b.append((
        "p",
        f"Mapping and inspection is the second use. Photogrammetry from RGB already builds impressive models {cite('nex2014','colomina2014')}. Depth or LiDAR is added when the operator needs metric scale, vegetation penetration, or a flight path that stays a fixed standoff from a girder {cite('nex2014','colomina2014','spencer2019','seo2018','dorafshan2018')}. Bridge work is the most documented civil case: drone-enabled visual inspection procedures exist, and unmanned systems have been used to cue fatigue-crack examination on steel members {cite('seo2018','dorafshan2018')}. Quality-assessment studies warn that a pretty mesh is not an inspection: motion blur, standoff, and lighting can hide the crack the flight was commissioned to find {cite('morgenthal2014')}. Next-best-view planners show how a ranging front-end can choose the next pose instead of flying a fixed lawnmower {cite('bircher2016')}. Civil-infrastructure reviews show that the bottleneck is rarely whether a point cloud can be built. It is converting that cloud into defect locations that an engineer will trust {cite('spencer2019')}. Aerial manipulation is a third, still-research use: a ranging front-end can keep a tool at a commanded standoff, but contact dynamics then dominate the stack {cite('ruggiero2018')}.",
    ))
    b.append(("fig", FIG / "fig3_uav_inspection_pipeline.png"))
    b.append((
        "caption",
        "Figure 5. Representative UAV inspection pipeline assembled from published practice. This figure is a literature synthesis, not a system designed in the present paper.",
    ))
    b.append((
        "p",
        f"Figure 5 describes a representative architecture assembled from published inspection practice, not a system designed in this paper. A bridge-inspection UAV typically carries a long-range ranging sensor (solid-state LiDAR or equivalent), a high-resolution RGB camera, and an IMU {cite('nex2014','colomina2014','spencer2019')}. Onboard software fuses range and inertia into a local metric volume and runs a lightweight network on RGB for component and surface-defect cues {cite('spencer2019')}. The aircraft uplinks an annotated model and geotagged defect hypotheses rather than raw multi-sensor video. The design is a bandwidth and trust design: inspectors re-photograph the flagged regions instead of watching hours of flight video. Multi-UAV structure-from-motion benefits from the same metric scale cue. Depth does not replace inter-vehicle communication, but it reduces the need for dense ground control {cite('nex2014','colomina2014')}.",
    ))

    b.append(("h2", "4.4 Collaborative robots and human-robot interaction"))
    b.append((
        "p",
        f"Collaborative cells fail first on safety, then on programming time {cite('villani2018','ajoudani2018','haddadin2017','lasota2017','robla2017')}. Depth cameras address both, with limits that must be stated. Broader HRI surveys make the same split: methods that keep a human safe, and methods that make the robot usable {cite('lasota2017','robla2017')}.",
    ))
    b.append((
        "p",
        f"Workspace monitoring. A ceiling or cell-side depth sensor builds a three-dimensional occupancy of the shared volume. Speed-and-separation monitoring, as framed by ISO/TS 15066, can then slow or stop the arm when a person enters a warning or protective zone {cite('iso15066','villani2018')}. The industrial-robot safety standard ISO 10218 still sits underneath that technical specification {cite('iso10218')}. Three-dimensional monitoring is more complete than a light curtain because it can shrink the protective volume as the arm slows and can see a person who leans over a fence. Control papers that quantify collaborative safety treat separation, speed, and a defined protective stop as metrics, not as a camera specification {cite('zanchettin2016')}. A consumer RGB-D camera is not automatically a safety-rated device. Certified implementations still need a safety controller, defined failure modes, and usually a redundant sensing channel {cite('villani2018','haddadin2017','iso15066','robla2017')}. Surveys of industrial human-robot collaboration treat safety and the user interface as the two bottlenecks that decide whether a cobot is actually used, not whether a depth demo exists {cite('villani2018','ajoudani2018','lasota2017')}. Collision detection on the arm itself remains necessary because vision cannot see inside the contact {cite('haddadin2017')}.",
    ))
    b.append((
        "p",
        f"Programming by demonstration. Tracking a tool or a marked workpiece in 3D lets a non-expert move the robot through a path {cite('villani2018','ajoudani2018')}. Structured light at short range is usually sufficient. The hard parts are correspondence, occlusion by the operator, and converting a human demonstration into a collision-free, force-feasible program. Gesture and activity cues can distinguish a small set of commands or detect that a station is ready for the next part {cite('shotton2011','villani2018')}. This is useful. It is not a substitute for a well-designed hardware enable switch.",
    ))

    b.append(("h2", "4.5 A copy-paste selection checklist"))
    b.append((
        "p",
        "The following list is intended to be copied into a requirements document. It is the operational summary of Sections 2-4.",
    ))
    b.append(("p", "1. Write the minimum and maximum range that must work, not the range that would be nice."))
    b.append(("p", "2. Write the worst lighting (sun through a door, welding flash, night) and the worst surface (black rubber, glass, wet steel)."))
    b.append(("p", "3. Write the mass, volume, and watt budget, including the illuminator."))
    b.append(("p", f"4. Write the safety integrity: visualization, planner input, or motor-stop. Motor-stop needs a certified channel {cite('iso15066','iso10218')}."))
    b.append(("p", "5. Pick the family from Figure 2. If two worlds appear, budget two sensors."))
    b.append(("p", f"6. Pick a SKU only after reading a peer-reviewed evaluation of that family, not a datasheet alone {cite('halmet2019','carfagni2019','tolgyessy2021')}."))
    b.append(("p", f"7. Budget calibration, time synchronization, and a warm-up interval as line items {cite('herrera2012','tolgyessy2021')}."))
    b.append(("p", f"8. Match the estimator to the cloud: TSDF for a table, VIO/VIDO for a room, LIO for a sparse outdoor scan {cite('newcombe2011','campos2021','xu2022fastlio')}."))
    b.append(("p", "9. Keep a raw or lightly filtered metric range path if the output can stop a motor. Do not stop a motor on a completed depth map alone."))
    b.append(("p", f"10. Re-test on the real surfaces and the real lighting. Indoor rankings reorder outdoors {cite('halmet2019')}."))

    b.append(("h1", "5. Challenges and future perspectives"))
    b.append(("h2", "5.1 Persistent deployment problems"))
    b.append((
        "p",
        "The research literature sometimes treats remaining errors as temporary. Fielded mechatronic systems suggest otherwise. Table 7 separates mitigations that already ship from ideas that remain research.",
    ))
    b.append((
        "p",
        f"Environment and materials. NIR-based cameras saturate in sunlight {cite('halmet2019','sarbolandi2015')}. Specular and transparent objects violate the single-bounce assumption and produce holes or biased ranges {cite('foix2011','kadambi2013')}. Coded and multi-frequency ToF can unmix some multipath, but they add capture time or hardware that a cheap module does not have {cite('kadambi2013','foix2011')}. Low-reflectance foam and black rubber starve the illuminator. Polarization, multi-frequency ToF, and radar fill-in help; none is universal. Comparative tests of ten indoor depth cameras remain the most useful reminder that material and lighting can reorder the ranking of otherwise similar devices {cite('halmet2019')}.",
    ))
    b.append((
        "p",
        f"Compute and power. Dense SLAM, TSDF fusion, and point-cloud networks are still expensive on a battery {cite('cadena2016','guo2021pc','gyongy2022')}. The illuminator, not only the GPU, shortens UAV endurance {cite('gyongy2022','floreano2015')}. Event-based and neuromorphic pipelines are promising where the scene is sparse in time {cite('gallego2022','falanga2020','lichtsteiner2008')}, but they require a different software stack.",
    ))
    b.append((
        "p",
        f"Calibration and time. Multi-sensor extrinsics drift with temperature and shock. A one-degree boresight error is negligible for a room-scale display and unacceptable for a 50 m LiDAR-camera fusion {cite('roriz2022','debeunne2020')}. Online self-calibration exists in research estimators {cite('qin2018','campos2021')}; it is not yet a maintenance-free commodity. Asynchronous streams and sensor dropouts can make a naively fused estimate worse than the best single sensor {cite('thrun2005','yeong2021')}.",
    ))
    b.append((
        "p",
        f"Interfaces and cost. Unlike USB webcams, depth devices still disagree on coordinate frames, distortion models, and confidence channels {cite('halmet2019','keselman2017')}. High-grade solid-state LiDAR remains expensive relative to the rest of a service robot {cite('roriz2022')}. Standardization would save more engineering time than another incremental denoiser.",
    ))
    b.append((
        "table",
        (
            "Table 7. Challenges and the honesty level of current mitigations.",
            ["Domain", "Concrete problem", "Pathway that already ships", "Pathway still in research"],
            [
                ["Environment", "Sunlight, glass, black parts", "dToF/LiDAR, multi-echo, radar", "Polarization / hyperspectral ToF"],
                ["Compute", "Dense 3D + semantics on-device", "NPU for 2D nets; sparse maps", f"Event + spiking co-design {cite('gallego2022','roy2019','davies2018')}"],
                ["Integration", "Extrinsics, time, dropouts", "Factory calibration; hardware sync", "Lifelong self-calibration"],
                ["Semantics", "Novel clutter, moving people", "Closed-set 3D detectors", "3D foundation models, Sim2Real"],
                ["Cost / API", "Proprietary frames, high unit cost", "ROS drivers; a few de-facto SDKs", "Safety-rated, interoperable depth"],
            ],
        ),
    ))

    b.append(("h2", "5.2 Outlook"))
    b.append((
        "p",
        f"Event-based and neuromorphic sensing. Asynchronous pixels report contrast changes at microsecond latency and high dynamic range {cite('lichtsteiner2008','brandli2014','gallego2022')}. Combined with pulsed illumination they can support low-power depth for high-speed robots {cite('falanga2020','he2024','zhou2021evo','rebecq2017','guo2024cmax')}. Automotive work has already shown that event cameras can close a low-latency perception loop that frame cameras miss {cite('gehrig2024')}. Co-design with spiking processors aims to cut the energy of spatial reasoning {cite('roy2019','davies2018')}. These systems will complement, not replace, frame-based RGB-D in the next product cycle.",
    ))
    b.append((
        "p",
        f"AI near the sensor. Depth completion and a first detection head are moving into image-signal processors and sensor stacks {cite('gyongy2022','eigen2014','ma2018')}. The engineering gain is bandwidth: a robot can ship a confidence-weighted cloud or a set of objects instead of a raw 30 Hz volume. The engineering risk is silent metric failure. Safety-related channels should keep a raw or lightly filtered range path.",
    ))
    b.append((
        "p",
        f"Richer fusion. Depth-RGB-IMU is already standard. Adding millimetre-wave radar (fog, rain, radial velocity) and thermal cameras is the practical next step, already visible in automotive datasets {cite('yeong2021','caesar2020')}. The unsolved part is a fusion integrity layer: knowing when to stop trusting a modality.",
    ))
    b.append((
        "p",
        f"Photonics and materials. Metalenses and other flat optics can shrink camera modules {cite('khorasaninejad2017')}. Chip-scale optical phased arrays and MEMS-on-photonics LiDAR attack the scanner itself {cite('sun2013opa','poulton2017','zhang2022mems','hsu2021opa')}. First-photon imaging and non-line-of-sight reconstructions show what is physically possible at extreme photon counts {cite('kirmani2014','otoole2018')}. They are not yet bill-of-materials options for a warehouse AMR. Neural radiance fields and 3D Gaussian splatting are changing how a reconstructed scene is stored and rendered {cite('mildenhall2020','kerbl2023')}. For a mechatronic stack they are still map representations, not replacements for a metric range channel that can stop a motor {cite('sucar2021','zhu2022nice')}.",
    ))
    b.append((
        "p",
        "Infrastructure sensing. Ceiling-mounted depth in factories and intelligent spaces can offload onboard mapping. The idea is sound. The obstacles are privacy, networking, and a shared metric frame that every vendor agrees to.",
    ))

    b.append(("h1", "6. Conclusion"))
    b.append((
        "p",
        "Depth vision in mechatronics did not evolve as a single replacement chain. Structured light made dense, inexpensive, close-range geometry ordinary. ToF made that geometry small enough for mobile heads and phones. Solid-state LiDAR made long-range ranging compatible with vibration-sensitive platforms. Algorithms (completion, tight-coupled odometry, RGB-D SLAM, and 3D semantics) converted those measurements into poses, maps, grasps, and safety zones.",
    ))
    b.append((
        "p",
        "The remaining limits are mostly physical and organizational: sunlight and materials, watts and memory, calibration, and fragmented interfaces. A next-generation system is unlikely to be one perfect depth camera. It is more likely to be a small set of complementary ranges, an estimator that knows its integrity, and, where possible, a shared map that the machine does not have to rebuild on every shift.",
    ))
    b.append((
        "p",
        "For practitioners, the operational conclusion is conservative. Choose the sensor by operating envelope, not by marketing generation. Budget calibration and lighting as first-class design items. Keep a metric, low-level range channel whenever the output can stop a motor. Treat public benchmarks as a map of where a method was tested, then re-test on the real surfaces, lighting, and safety integrity of the cell or airframe.",
    ))

    b.append(("h1", "Author contributions"))
    b.append((
        "p",
        "Bo Zhang drafted the review, organized the literature, and prepared the comparative tables and figures. Hongchao Cui conceived the hardware-centered scope, revised the manuscript, and supervised the work. Both authors approved the final version.",
    ))
    b.append(("h1", "Funding"))
    b.append(("p", "This review received no specific external funding. The authors will replace this sentence with grant numbers if required by their institution."))
    b.append(("h1", "Conflicts of interest"))
    b.append(("p", "The authors declare no conflict of interest."))
    b.append(("h1", "Data availability"))
    b.append(("p", "No new datasets were generated. All cited sources are listed in the References."))
    b.append(("h1", "Acknowledgments"))
    b.append(("p", "The authors thank colleagues who commented on an earlier draft. Any remaining errors are the authors' own."))

    b.append(("refs", None))
    return b


def blocks_to_docx(blocks: list[tuple], path: Path) -> None:
    doc = Document()
    style_doc(doc)
    for kind, payload in blocks:
        if kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            add_text(p, payload, size=16, bold=True)
        elif kind == "authors":
            para(doc, payload, first_line=False, center=True, size=12, space_after=4)
        elif kind == "affil":
            for line in payload.split("\n"):
                para(doc, line, first_line=False, center=True, size=10, italic=True, space_after=2)
        elif kind == "corr":
            para(doc, payload, first_line=False, center=True, size=10, space_after=16)
        elif kind == "h1":
            heading(doc, payload, 1)
        elif kind == "h2":
            heading(doc, payload, 2)
        elif kind == "p":
            if payload.startswith("Keywords:"):
                para(doc, payload, first_line=False, italic=True, size=10)
            elif payload.startswith("•") or (len(payload) > 2 and payload[0].isdigit() and payload[1] in ".)"):
                para(doc, payload, first_line=False, size=11)
            else:
                para(doc, payload)
        elif kind == "caption":
            caption(doc, payload)
        elif kind == "fig":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.add_run().add_picture(str(payload), width=Inches(6.3))
        elif kind == "table":
            title, headers, rows = payload
            add_table(doc, headers, rows, title)
        elif kind == "refs":
            heading(doc, "References", 1)
            for i, key in enumerate(_order, 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.35)
                p.paragraph_format.first_line_indent = Inches(-0.35)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                add_text(p, f"[{i}] {BIB[key]}", size=10)
    doc.save(path)


def main() -> None:
    blocks = manuscript_blocks()
    unused = sorted(set(BIB) - set(_order))
    md_path = ROOT / "Zhang_Cui_Depth_Vision_Sensors_Review_revised.md"
    docx_path = ROOT / "Zhang_Cui_Depth_Vision_Sensors_Review_revised.docx"
    write_markdown(md_path, blocks)
    blocks_to_docx(blocks, docx_path)
    words = sum(len(p.split()) for k, p in blocks if k == "p" and isinstance(p, str))
    print(f"references cited: {len(_order)}")
    print(f"references unused: {len(unused)} -> {unused}")
    print(f"body word count (paragraphs only): {words}")
    print(f"wrote {md_path}")
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    main()
