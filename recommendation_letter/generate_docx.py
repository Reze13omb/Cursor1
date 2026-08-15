#!/usr/bin/env python3
"""Generate Word versions of the recommendation letter package."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent


def set_run_font(run, name_en="Times New Roman", name_zh="Songti SC", size=12, bold=False):
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name_zh)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc, text, *, size=12, bold=False, space_after=8, first_line=True, align="left"):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if first_line:
        pf.first_line_indent = Inches(0.3)
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf.first_line_indent = Inches(0)
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Inches(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def new_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    return doc


def write_revised_en():
    doc = new_doc()
    add_para(doc, "August 12, 2026", first_line=False, space_after=18)
    add_para(doc, "To the Admissions Committee:", first_line=False, space_after=12)

    body = [
        "I am Yan Qi, an analyst in the Equity Research Department of Minmetals Securities Co., Ltd., Beijing Branch. Bo Zhang completed a six-month internship in our department from January to June 2026. I was his direct supervisor throughout that period and worked with him closely enough to observe his daily progress, working habits, learning methods, and professional judgment. I write to offer my strong recommendation of Bo Zhang for admission to your graduate program.",
        "When Bo first joined the team, he was still learning our research workflow, as most new student interns are. He distinguished himself quickly. Even the early assignments I gave him were treated with unusual care: he checked each item thoroughly, and the materials he submitted were well organized and nearly error-free. That reliability reduced the time our team spent on the preparatory work that precedes financial analysis, and it is the reason I was able to give him greater responsibility.",
        "Much of the commercial and financial content was new to him, so the learning demand was high. He did not wait to be walked through the material. He took the initiative to study the business positioning and objectives, our service details, and the competitive environment. When he encountered a problem he could not resolve on his own, he came to me with structured notes, a clear account of what he already understood, and specific questions. As I came to know his working style, I also saw that he would use the data he had assembled to form his own preliminary views and then discuss them with me. Those ideas were still developing, but they showed sound logical thinking and a genuine willingness to go deeper than the assigned task required.",
        "In the workplace, Bo has been highly reliable, intellectually curious, and strongly task-driven, with clear room to grow. I have often encouraged him to pursue a master’s degree so that he can systematically strengthen his data-analysis and research skills, connect theoretical training with the kind of industry problems he encountered here, and prepare himself for more advanced professional work. Graduate study is, in my view, the right next step for him.",
        "I recommend Bo Zhang to your program with confidence. His diligence, intellectual honesty, and ability to learn quickly from unfamiliar material give me every reason to believe that he will adapt well to the pace of graduate study and continue to improve as a researcher. I would be glad to provide any further information the committee may need.",
    ]
    for para in body:
        add_para(doc, para, space_after=10)

    add_para(doc, "Sincerely,", first_line=False, space_after=36)
    for line in [
        "Yan Qi",
        "Analyst",
        "Equity Research Department",
        "Minmetals Securities Co., Ltd., Beijing Branch",
        "Email: qiyan1@wkzq.com.cn",
        "Tel: +86-10-56307033",
        "Address: Room 301 (Units 02, 03B), 3rd Floor, No. 7 North Chaoyangmen Avenue, Dongcheng District, Beijing, China",
    ]:
        add_para(doc, line, first_line=False, space_after=0, size=12)

    path = OUT_DIR / "RL_Bo_Zhang_Yan_Qi_revised.docx"
    doc.save(path)
    return path


def write_package():
    doc = new_doc()
    add_para(doc, "Recommendation Letter Package for Bo Zhang", size=16, bold=True, first_line=False, align="center", space_after=6)
    add_para(doc, "Recommender: Yan Qi, Minmetals Securities", size=12, first_line=False, align="center", space_after=18)

    add_para(doc, "一、原文忠实译文", size=14, bold=True, first_line=False, space_after=10)
    translation = [
        "2026年8月12日",
        "尊敬的先生／女士：",
        "我是齐岩，就职于五矿证券有限公司北京分公司股权研究部，担任分析师。张博于今年1月至6月在我部门完成实习。在此期间，我是他的直接主管，密切关注其日常工作进展，并观察了他的工作习惯、学习方法及人际沟通能力。现特此为他的研究生申请撰写这封有力的推荐信。",
        "他刚到部门实习时，和大多数新来的学生实习生一样，对工作流程尚不熟悉。但他很快便在其他实习生中脱颖而出。尽管起初我交给他的大多是基础性工作，他都以认真的态度对待，并会逐项仔细核对。他提交的材料条理清晰、差错极少，大大减轻了团队在财务分析前期准备工作中的负担。",
        "此后，凭借积极的态度，他获得了参与新业务板块开发的机会。其中大部分业务与财务知识对他而言都是全新的，需要他投入更多精力去理解。他主动研究业务定位与目标、我们的服务细节以及竞争环境。遇到自己想不明白的问题，他会带着整理好的思考笔记来与我沟通。随着对他逐渐熟悉，我发现他会有意识地在整理好的数据基础上提出自己虽尚显初步、但有独立见解的想法，这表明他具备良好的逻辑思维能力，并愿意深入探究。",
        "根据我在工作中对张博的观察，他非常可靠、求知欲强、执行力出色，并具有相当大的成长潜力。我常鼓励他攻读硕士学位，以系统提升数据分析与研究能力，将理论知识与真实行业案例相结合，拓宽国际视野，成长为具有国际视野的跨学科专业人才。",
        "我诚挚地向贵校推荐张博。我相信，凭借他勤奋的品质，他能够很好地适应研究生阶段的学习节奏，并在学术研究中持续提升自己。",
        "此致",
        "齐岩",
        "分析师",
        "股权研究部（待确认）",
        "五矿证券有限公司北京分公司",
        "邮箱：qiyan1@wkzq.com.cn",
        "电话：86-10-56307033",
        "地址：中国北京市东城区朝阳门北大街7号3层301室（02、03B单元）",
    ]
    for i, para in enumerate(translation):
        first = i not in (0, 1) and not para.startswith(("此致", "齐岩", "分析师", "股权", "五矿", "邮箱", "电话", "地址"))
        add_para(doc, para, first_line=first, space_after=8)

    add_para(doc, "二、修订后英文稿（建议提交此版）", size=14, bold=True, first_line=False, space_after=10)
    add_para(
        doc,
        "Please see the standalone file RL_Bo_Zhang_Yan_Qi_revised.docx for a clean submission copy. The revised English text is also included below.",
        first_line=False,
        space_after=10,
        size=11,
    )

    en_paras = [
        "August 12, 2026",
        "To the Admissions Committee:",
        "I am Yan Qi, an analyst in the Equity Research Department of Minmetals Securities Co., Ltd., Beijing Branch. Bo Zhang completed a six-month internship in our department from January to June 2026. I was his direct supervisor throughout that period and worked with him closely enough to observe his daily progress, working habits, learning methods, and professional judgment. I write to offer my strong recommendation of Bo Zhang for admission to your graduate program.",
        "When Bo first joined the team, he was still learning our research workflow, as most new student interns are. He distinguished himself quickly. Even the early assignments I gave him were treated with unusual care: he checked each item thoroughly, and the materials he submitted were well organized and nearly error-free. That reliability reduced the time our team spent on the preparatory work that precedes financial analysis, and it is the reason I was able to give him greater responsibility.",
        "Much of the commercial and financial content was new to him, so the learning demand was high. He did not wait to be walked through the material. He took the initiative to study the business positioning and objectives, our service details, and the competitive environment. When he encountered a problem he could not resolve on his own, he came to me with structured notes, a clear account of what he already understood, and specific questions. As I came to know his working style, I also saw that he would use the data he had assembled to form his own preliminary views and then discuss them with me. Those ideas were still developing, but they showed sound logical thinking and a genuine willingness to go deeper than the assigned task required.",
        "In the workplace, Bo has been highly reliable, intellectually curious, and strongly task-driven, with clear room to grow. I have often encouraged him to pursue a master’s degree so that he can systematically strengthen his data-analysis and research skills, connect theoretical training with the kind of industry problems he encountered here, and prepare himself for more advanced professional work. Graduate study is, in my view, the right next step for him.",
        "I recommend Bo Zhang to your program with confidence. His diligence, intellectual honesty, and ability to learn quickly from unfamiliar material give me every reason to believe that he will adapt well to the pace of graduate study and continue to improve as a researcher. I would be glad to provide any further information the committee may need.",
        "Sincerely,",
        "Yan Qi",
        "Analyst",
        "Equity Research Department",
        "Minmetals Securities Co., Ltd., Beijing Branch",
        "Email: qiyan1@wkzq.com.cn",
        "Tel: +86-10-56307033",
        "Address: Room 301 (Units 02, 03B), 3rd Floor, No. 7 North Chaoyangmen Avenue, Dongcheng District, Beijing, China",
    ]
    for i, para in enumerate(en_paras):
        first = i >= 2 and i <= 6
        add_para(doc, para, first_line=first, space_after=8)

    add_para(doc, "三、修订稿中文对照（供推荐人审阅）", size=14, bold=True, first_line=False, space_after=10)
    zh_revised = [
        "2026年8月12日",
        "致招生委员会：",
        "我是齐岩，五矿证券有限公司北京分公司股权研究部分析师。张博于2026年1月至6月在我部门完成为期六个月的实习。在此期间，我是他的直接主管，日常接触足以观察他的工作进展、工作习惯、学习方法及专业判断。现特此为他申请贵校研究生项目提供有力推荐。",
        "张博刚加入团队时，和大多数新来的学生实习生一样，仍在熟悉我们的研究工作流程。但他很快便脱颖而出。即便是我最初交给他的任务，他也处理得格外认真：逐项仔细核对，提交的材料条理清晰、几乎没有差错。这种可靠性减轻了团队在财务分析前期准备上的时间投入，也因此我能够把更重要的工作交给他。",
        "此后，大量业务与财务内容对他而言都是新的，学习压力不小。他没有等待别人带着走，而是主动研究业务定位与目标、我们的服务细节以及竞争环境。遇到自己无法独立解决的问题，他会带着结构清晰的笔记、对自己已知内容的说明，以及具体问题来找我。随着我逐渐了解他的工作方式，我也看到他会在已整理的数据基础上形成自己的初步判断，再与我讨论。这些想法当时仍在发展中，但已经体现出扎实的逻辑思维，以及愿意把问题想得比任务本身更深一层的态度。",
        "在工作中，张博一贯可靠、求知欲强、执行力突出，并有明确的成长空间。我常鼓励他攻读硕士，以便系统加强数据分析与研究能力，把理论训练与他在这里接触到的行业问题结合起来，并为更进一步的专业工作做好准备。在我看来，研究生阶段是他合适的下一步。",
        "我有信心向贵项目推荐张博。他的勤奋、治学诚实，以及从陌生材料中快速学习的能力，使我有充分理由相信：他能够适应研究生阶段的学习节奏，并在研究中持续进步。如委员会需要更多信息，我乐意补充。",
        "此致",
        "齐岩",
        "分析师",
        "股权研究部",
        "五矿证券有限公司北京分公司",
    ]
    for i, para in enumerate(zh_revised):
        first = 2 <= i <= 6
        add_para(doc, para, first_line=first, space_after=8)

    path = OUT_DIR / "RL_Bo_Zhang_package_zh_en.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    p1 = write_revised_en()
    p2 = write_package()
    print(p1)
    print(p2)
