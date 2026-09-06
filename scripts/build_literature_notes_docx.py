#!/usr/bin/env python3
"""Build Week 1–3 literature notes as a .docx."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc, text, *, bold=False, italic=False, size=12, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_mixed(doc, parts, *, space_after=6, space_before=0, size=12):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.15
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_h1(doc, text):
    return add_para(doc, text, bold=True, size=16, space_before=4, space_after=10, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_h2(doc, text):
    return add_para(doc, text, bold=True, size=14, space_before=14, space_after=8)


def add_h3(doc, text):
    return add_para(doc, text, bold=True, size=12, space_before=10, space_after=4)


def add_bullet(doc, lead, text=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if lead:
        r1 = p.add_run(lead)
        set_run_font(r1, bold=True)
    if text:
        r2 = p.add_run(text)
        set_run_font(r2)
    return p


def set_cell(cell, text, *, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def paper_block(doc, title, cite, one_liner, problem, method, results, limits, use):
    add_h3(doc, title)
    add_mixed(doc, [("完整引用：", True, False), (cite, False, False)])
    add_mixed(doc, [("一句话：", True, False), (one_liner, False, False)])
    add_mixed(doc, [("要解决的问题：", True, False), (problem, False, False)])
    add_mixed(doc, [("方法：", True, False), (method, False, False)])
    add_mixed(doc, [("关键结论：", True, False), (results, False, False)])
    add_mixed(doc, [("局限：", True, False), (limits, False, False)])
    add_mixed(doc, [("对本课题的用处：", True, False), (use, False, False)])


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    add_h1(doc, "Literature Notes / 文献笔记")
    add_para(
        doc,
        "Week 1–3  |  Gravity Compensation of a Planar Five-Bar Mechanism Using Gear-Spring Modules",
        bold=True,
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_para(
        doc,
        "Student: Bo Zhang (8571260)    Supervisor: Dr Chin-Hsing Kuo    Date: September 2026",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    add_h2(doc, "0. 怎么读、这周要记住什么")
    add_para(
        doc,
        "本笔记只服务 Week 1–3：把静平衡、GSM、Delta 方法和平面五杆运动学读清楚。每篇只记问题、方法、结论、局限、和本课题的关系。GSM 弹簧力矩公式、k/ψ 设计和仿真放到 Week 4–6。",
    )
    add_para(doc, "组会上用三句话交代文献位置：", space_after=4)
    add_bullet(doc, "本课题不复现 Delta。", " 2020 年 Delta+GSM 论文是方法模板：目标构型近似 + 普通压缩弹簧 + 力矩/能耗指标。")
    add_bullet(doc, "GSM 本身来自串联臂论文。", " 本课题要把它装到平面闭环五杆的两个驱动关节上，静力学必须重推。")
    add_bullet(doc, "五杆文献提供逆解、奇异和工作空间。", " 不提供重力补偿；补偿公式要自己写。")

    add_h2(doc, "1. 方法模板：Delta + GSM（必读，最重要）")
    paper_block(
        doc,
        "1.1 Nguyen, Lin and Kuo (2020), Mechanism and Machine Theory",
        "V. L. Nguyen, C.-Y. Lin and C.-H. Kuo, “Gravity compensation design of Delta parallel robots using gear-spring modules,” Mechanism and Machine Theory, vol. 154, 104046, 2020. DOI: 10.1016/j.mechmachtheory.2020.104046",
        "给 Delta 每条近端臂装一个 GSM，在对称目标构型上解析逼近完美静平衡，再用 TRR / GCD / ERR 评价。",
        "并联机器人重力补偿难：腿和动平台耦合。配重会增惯量；很多弹簧方案要理想零自由长度弹簧（ZFL spring），安装和摩擦麻烦。作者要把已在串联臂上验证的 GSM 做到 Delta 上，且不牺牲工作空间。",
        "三条近端臂各装一个 GSM（齿轮1固连机座，齿轮2随臂转，压缩弹簧经连杆和齿轮出补偿力矩）。全工作空间完美平衡做不到，因此选对称目标构型（三腿同角、动平台在竖直轴上），令 Ts ≈ Tw，解析求刚度 k 和安装角 ψ。齿轮比取 ng = 2，是为了让简化后的弹簧力矩形式和目标构型上的重力力矩对齐。指标：TRR、M-TRR、P-TRR、GCD、ERR。",
        "理论模型：约 86% 工作空间力矩下降为正（GCD，阈值 0）。FANUC M-3iA/12H 估计模型、10 kg 负载、取放任务：峰值力矩约降 38.4%，能耗约降 55.4%。对比纯扭转弹簧和 ZFL 拉伸弹簧，GSM 明显更好。目标角按任务不同，最优值不同（取放常选 0，螺旋运动 π/6 更好）。",
        "不是全空间完美平衡；设计阶段忽略摩擦和惯性；FANUC 是估计尺寸的数值例子，没有空间样机；负载变化时 k、ψ 要重调。",
        "本课题直接沿用：目标构型近似、普通压缩弹簧、TRR 一类指标、ng=2 需要重新检查。必须改写的：对称条件（五杆是中线 θL = θR = Θ，不是三腿同角）、Tw 公式、两套而非三套 GSM。Week 4 才开始套 GSM 公式，这周先吃透“为什么只能在目标构型上完美”。",
    )

    add_h2(doc, "2. GSM 从哪来：串联臂论文（必读）")
    paper_block(
        doc,
        "2.1 Nguyen, Lin and Kuo (2020), ASME Journal of Mechanisms and Robotics",
        "V. L. Nguyen, C.-Y. Lin and C.-H. Kuo, “Gravity compensation design of planar articulated robotic arms using the gear-spring modules,” ASME J. Mechanisms and Robotics, vol. 12, no. 3, 031014, 2020. DOI: 10.1115/1.4045650",
        "提出 GSM：带压缩弹簧的齿轮–滑块模块，装在平面串联臂关节上做重力补偿。",
        "串联臂重力补偿若用平行四边形或 ZFL 弹簧，往往体积大、难模块化。需要一个紧凑、用普通弹簧、可装到关节上的模块。",
        "每个关节一个 GSM。弹簧刚度可用优化，也可用解析逼近完美平衡。文中比较 1、2、3 自由度臂，并考虑齿轮接触功率损失，给出刚度修正。单自由度样机实验验证。",
        "解析逼近和优化结果接近。单自由度实验：装 GSM 后电机功耗约降 86.5%。",
        "对象是开链串联臂，不是闭环并联。实验也是单自由度台架，不是完整多腿并联机构。",
        "本课题的模块几何、弹簧力矩思路、d0 = 0 和齿轮臂远短于连杆等简化假设，都从这里来。五杆是闭环，Tw 不能按串联臂逐关节拆，这是本课题要新推的部分。",
    )

    add_h2(doc, "3. 静平衡总图：综述与经典（必读背景）")
    paper_block(
        doc,
        "3.1 Arakelian (2016), Advanced Robotics",
        "V. Arakelian, “Gravity compensation in robotics,” Advanced Robotics, vol. 30, no. 2, pp. 79–96, 2016. DOI: 10.1080/01691864.2015.1090334",
        "按补偿力的来源把机器人重力补偿分成三类：配重、弹簧、辅助驱动，再按结构细分。",
        "关节要长期抵抗连杆自重，电机负担大。文献方法很多，需要一张能分类比较的地图。",
        "综述。配重：让系统质心不动，但惯量增加。弹簧：分 ZFL 与非 ZFL；ZFL 便于完美平衡，但实物常要用滑轮、导轨或预紧来近似。主动补偿：另加执行器，不是纯被动。",
        "被动弹簧方案在能耗和安全性上通常优于配重；完美平衡往往依赖 ZFL 或特殊几何；变负载仍是难点。",
        "综述不给可算的设计公式；对具体机构要另找论文。",
        "写绪论和文献综述时用这篇分类。本课题属于“弹簧 + 辅助机构（齿轮滑块）+ 非 ZFL 普通压缩弹簧”。向老师解释时：我们不做配重，也不做电子主动补偿。",
    )
    paper_block(
        doc,
        "3.2 Herder (2001), PhD thesis, TU Delft",
        "J. L. Herder, Energy-free Systems: Theory, Conception and Design of Statically Balanced Spring Mechanisms, Ph.D. thesis, Delft University of Technology, 2001.",
        "静平衡的理论源头：用弹簧势能抵消重力势能，使系统在工作范围内近似“零刚度、零重力力矩”。",
        "如何从能量观点设计弹簧平衡机构，而不是只在一个姿态配平。",
        "完美静平衡：任意位形重力势能变化都被弹簧势能抵消，驱动关节重力力矩恒为零。大量讨论 ZFL 弹簧、储能元件布置和能量自由调节。",
        "给出后来几乎所有 gravity equilibrator 论文都引用的定义和设计哲学。",
        "是理论专著，不是某个机器人的现成配方；本科不必通读全书，读清“完美平衡”定义即可。",
        "本课题的评价标准要诚实：我们做的是目标构型上的近似平衡，不是 Herder 意义上的全工作空间完美平衡。测力计读数不会到零，剩余力来自近似误差和摩擦。",
    )

    add_h2(doc, "4. 平面五杆本身：运动学（Week 1–3 推导要用）")
    paper_block(
        doc,
        "4.1 Liu, Wang and Pritschow (2006), Mechanism and Machine Theory",
        "X.-J. Liu, J. Wang and G. Pritschow, “Kinematics, singularity and workspace of planar 5R symmetrical parallel mechanisms,” Mechanism and Machine Theory, vol. 41, no. 2, pp. 145–169, 2006. DOI: 10.1016/j.mechmachtheory.2005.05.004",
        "系统给出对称平面 5R（五杆）并联机构的运动学、奇异和工作空间。",
        "五杆看起来简单，但逆解有多组、有奇异、工作空间形状随杆长比大变，设计时必须先分清可用区域。",
        "对称 5R：机架 + 两驱动杆 + 两连杆，末端在两连杆铰点，2 自由度。给出正逆解、奇异分类（一般含逆解奇异 / 正解奇异）和工作空间边界。同组还有性能图谱与尺度综合的姊妹篇（同卷 pp. 119–144）。",
        "对称五杆的工作空间关于中线对称；奇异会把工作空间割成几块，轨迹不能随便穿过。",
        "这篇只做运动学和尺度，不做重力补偿，也没有弹簧。",
        "Week 1–3 写逆解和标奇异时以这篇为骨架。本课题布置是固定杆在上、两杆下垂，坐标原点和重力方向要自己定，不能照抄他们“机架在下”的图。选尺寸时避开奇异密集区，并保证中线提升路径在可达工作空间内。",
    )

    add_h2(doc, "5. 对照：别人怎么给 Delta / 经典机构做弹簧平衡")
    paper_block(
        doc,
        "5.1 Simionescu, Ciupitu and Ionita (2015), Mechanism and Machine Theory",
        "I. Simionescu, L. Ciupitu and L. C. Ionita, “Static balancing with elastic systems of DELTA parallel robots,” Mechanism and Machine Theory, vol. 87, pp. 150–162, 2015. DOI: 10.1016/j.mechmachtheory.2014.11.008",
        "用弹性系统给 Delta 做静平衡，给出完美平衡（更多 ZFL 弹簧 + 辅助机构）和近似平衡（每腿一根 ZFL 弹簧）两套思路。",
        "Delta 要同时平衡三条腿和动平台。如何用弹簧做到、要几根弹簧、要不要改近端臂结构。",
        "完美方案：近端臂做成平行四边形并用滑块摇杆约束，每腿多根 ZFL 弹簧。近似方案：每腿一根 ZFL 弹簧挂在近端臂上，结构简单但不是完美平衡。",
        "说明 Delta 的“完美”往往以辅助机构和 ZFL 为代价；近似方案更接近工程。",
        "依赖 ZFL；辅助机构占空间。Nguyen et al. (2020) 正是针对这些缺点提出 GSM。",
        "写文献对比时用：本课题和 2015 这篇、2020 GSM 篇走的是“近似 + 普通弹簧 + 紧凑模块”，不是再做一套 ZFL 完美平衡。",
    )
    paper_block(
        doc,
        "5.2 French and Widden (2000), Proc. IMechE Part C",
        "M. J. French and M. B. Widden, “The spring-and-lever balancing mechanism, George Carwardine and the Anglepoise lamp,” Proc. IMechE, Part C: Journal of Mechanical Engineering Science, vol. 214, no. 3, pp. 501–508, 2000. DOI: 10.1243/0954406001523137",
        "经典弹簧杠杆平衡（Anglepoise 台灯）的简明分析：自由长度近似为零的弹簧，按合适刚度可做到完美平衡。",
        "从最简单的弹簧–杠杆把“为什么弹簧能抵消重力”讲清楚，并追溯 Carwardine 的专利。",
        "对单自由度和两自由度台灯式机构做精确分析，并计入构件自重。",
        "完美平衡的几何直觉：弹簧力臂和重力力臂随角度一起变，才能处处抵消。",
        "对象是开式台灯臂，不是闭环五杆；用的是接近 ZFL 的弹簧模型。",
        "建立直觉用，不当本课题主方法。若组会有人问“弹簧为什么能平衡”，用这篇比用 Delta 公式更合适。",
    )

    add_h2(doc, "6. 对照总表（写综述和开题答辩用）")
    table = doc.add_table(rows=8, cols=4)
    table.style = "Table Grid"
    header = ["文献", "对象", "补偿手段", "和本课题的关系"]
    rows = [
        header,
        ["Nguyen et al., 2020 MMT", "Delta 并联", "每腿一个 GSM，目标构型近似", "方法模板；对象要换成五杆"],
        ["Nguyen et al., 2020 JMR", "平面串联臂", "关节 GSM，解析/优化定 k", "模块来源；闭环力矩要重推"],
        ["Arakelian, 2016", "综述", "配重 / 弹簧 / 主动", "绪论分类；本课题属非 ZFL 弹簧"],
        ["Herder, 2001", "一般弹簧机构", "能量等价、完美平衡", "定义标准；我们做的是近似"],
        ["Liu et al., 2006", "对称平面 5R", "无补偿，只做运动学", "逆解、奇异、工作空间"],
        ["Simionescu et al., 2015", "Delta", "ZFL 弹簧，完美/近似", "反衬：我们不用 ZFL"],
        ["French & Widden, 2000", "台灯式杠杆", "接近 ZFL 的弹簧杠杆", "直觉；非主方法"],
    ]
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            set_cell(table.rows[i].cells[j], txt, bold=(i == 0), size=10)
    doc.add_paragraph()

    add_h2(doc, "7. 文献缺口（对应本课题要做的事）")
    add_bullet(doc, "GSM 还没做到平面闭环五杆上。", " 串联和 Delta 都有了，上置机架、两腿下垂的五杆没有现成 Tw 和目标构型定义。")
    add_bullet(doc, "目标构型仍靠手工试几个角。", " Delta 论文已显示任务不同最优角不同，但没有按轨迹自动选。这是本课题第 6 章。")
    add_bullet(doc, "Delta+GSM 没有做成可测样机。", " 五杆可以打印；验证用手持机械测力计（IMADA PS-10N）测移动所需的力，不用电子感应。")
    add_para(
        doc,
        "Week 1–3 文献读完后，下一步不是再找论文，而是：定五杆尺寸草图，写逆运动学，推导 Tw。",
        space_before=6,
    )

    add_h2(doc, "8. 组会 2 分钟口述（可直接念）")
    add_para(
        doc,
        "我这三周的文献主线是重力补偿。Arakelian 把方法分成配重、弹簧和主动补偿；Herder 把完美静平衡定义成势能处处抵消。Kuo 老师组的 GSM 用普通压缩弹簧做成紧凑模块，先用在串联臂，再用在 Delta：因为全空间完美平衡做不到，所以在对称目标构型上解析匹配力矩。Delta 那篇没有样机，五杆运动学可以参考 Liu 的 5R 论文，但重力力矩和目标构型都要重推。我的毕设就是把这套方法换到上置机架的平面五杆上，并加一章按轨迹选目标角；实验用手持机械测力计，不碰感应。",
    )

    add_h2(doc, "9. 起步参考文献（可直接贴进论文）")
    refs = [
        "V. L. Nguyen, C.-Y. Lin and C.-H. Kuo, “Gravity compensation design of Delta parallel robots using gear-spring modules,” Mechanism and Machine Theory, vol. 154, 104046, 2020.",
        "V. L. Nguyen, C.-Y. Lin and C.-H. Kuo, “Gravity compensation design of planar articulated robotic arms using the gear-spring modules,” ASME Journal of Mechanisms and Robotics, vol. 12, no. 3, 031014, 2020.",
        "V. Arakelian, “Gravity compensation in robotics,” Advanced Robotics, vol. 30, no. 2, pp. 79–96, 2016.",
        "J. L. Herder, Energy-free Systems: Theory, Conception and Design of Statically Balanced Spring Mechanisms, Ph.D. thesis, Delft University of Technology, 2001.",
        "X.-J. Liu, J. Wang and G. Pritschow, “Kinematics, singularity and workspace of planar 5R symmetrical parallel mechanisms,” Mechanism and Machine Theory, vol. 41, no. 2, pp. 145–169, 2006.",
        "X.-J. Liu, J. Wang and G. Pritschow, “Performance atlases and optimum design of planar 5R symmetrical parallel mechanisms,” Mechanism and Machine Theory, vol. 41, no. 2, pp. 119–144, 2006.",
        "I. Simionescu, L. Ciupitu and L. C. Ionita, “Static balancing with elastic systems of DELTA parallel robots,” Mechanism and Machine Theory, vol. 87, pp. 150–162, 2015.",
        "M. J. French and M. B. Widden, “The spring-and-lever balancing mechanism, George Carwardine and the Anglepoise lamp,” Proc. IMechE, Part C, vol. 214, no. 3, pp. 501–508, 2000.",
        "R. Clavel, “A fast robot with parallel geometry,” Proc. Int. Symp. Industrial Robots, pp. 91–100, 1988.",
    ]
    for i, ref in enumerate(refs, 1):
        add_para(doc, f"[{i}] {ref}", space_after=3)

    out = Path("/workspace/docs/Literature_Notes_Week1-3_Bo_Zhang.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    artifact = Path("/opt/cursor/artifacts")
    if artifact.is_dir():
        doc.save(artifact / "Literature_Notes_Week1-3_Bo_Zhang.docx")
    print("wrote", out)


if __name__ == "__main__":
    build()
