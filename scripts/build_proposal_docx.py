#!/usr/bin/env python3
"""Build the revised English thesis proposal as a .docx."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc, text, *, bold=False, italic=False, size=12, space_after=8, space_before=0, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_mixed(doc, parts, *, space_after=8, space_before=0):
    """parts: list of (text, bold, italic)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_run_font(run, bold=bold, italic=italic)
    return p


def add_heading_num(doc, text):
    return add_para(doc, text, bold=True, size=14, space_before=14, space_after=8)


def add_subhead(doc, text):
    return add_para(doc, text, bold=True, size=12, space_before=10, space_after=6)


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if bold_lead:
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def set_cell_text(cell, text, *, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=bold)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_para(
        doc,
        "Undergraduate Thesis Project Proposal",
        bold=True,
        size=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    add_para(
        doc,
        "Gravity Compensation of a Planar Five-Bar Mechanism Using Gear-Spring Modules: Analytical Design, Task-Oriented Optimization, and Prototype Validation",
        bold=True,
        size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )
    add_para(
        doc,
        "Student: Bo Zhang    Student number: 8571260\nSupervisor: Dr Chin-Hsing Kuo    Date: August 2026",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=14,
    )

    add_heading_num(doc, "1. Abstract")
    add_para(
        doc,
        "This project develops a gravity-compensation design for a planar closed-loop five-bar mechanism using gear-spring modules (GSMs). The work follows the analytical approximation method established for Delta parallel robots by Nguyen, Lin and Kuo (2020), but does not reproduce that robot. The five-bar is chosen because it keeps the same design logic—compact modules, practical compression springs, and balancing at targeted configurations—while remaining simple enough for complete theoretical derivation, numerical simulation, and a 3D-printed prototype within an undergraduate thesis.",
    )
    add_para(
        doc,
        "Two actuated cranks will each carry one GSM. Spring stiffness and installation angle will be determined so that spring torque approximately cancels gravitational torque at a set of symmetric targeted configurations. A dedicated chapter will then treat the targeted configuration as a design variable and select it automatically for given planar trajectories, rather than choosing it by hand.",
    )
    add_para(
        doc,
        "Prototype validation will not use electronic force sensing, load cells, motor-current estimation, or any other sensing technology. The design will be checked by a simpler mechanical principle: the residual unbalance of the mechanism appears as the external force required to move it slowly by hand. That force will be measured with a handheld analog mechanical force gauge (for example an IMADA PS-10N, 10 N capacity, 0.05 N resolution), comparing the GSM-on and GSM-off conditions along the same paths.",
    )
    add_para(
        doc,
        "The expected outcome is a complete theory–simulation–experiment study, and a clear design rule for task-oriented selection of the targeted configuration on a planar five-bar.",
    )

    add_heading_num(doc, "2. Background and motivation")
    add_para(
        doc,
        "Gravity compensation (static balancing) reduces the motor torque needed to support a robot’s own weight. Springs or counterweights store and release energy so that gravitational potential varies less over the workspace. This improves energy efficiency, allows smaller actuators, and can make physical human–robot interaction safer.",
    )
    add_para(
        doc,
        "The GSM is a compact geared five-bar / gear-slider module with a practical compression spring. It was first shown on planar serial arms, then applied to Delta parallel robots (Nguyen et al., 2020). For the Delta, three identical GSMs are mounted on the proximal links. Perfect balancing over the whole workspace is not achievable, so the design enforces zero residual gravitational torque only at symmetric targeted configurations, and accepts an approximation nearby. Torque- and energy-reduction indices are then used to evaluate performance.",
    )
    add_para(
        doc,
        "A planar five-bar with five revolute joints has two degrees of freedom. In the layout used here, the ground link is fixed at the top and two cranks hang and swing beneath it. The remaining two links close the loop; the end-effector sits at the coupler joint. This is kinematically simpler than a Delta (two actuated legs instead of three, planar instead of spatial), yet gravity compensation is still non-trivial because the two legs are coupled. The same GSM idea can be installed on the two actuated joints. Because the five-bar does not have the Delta’s three-fold symmetry, the gravitational-torque model, the definition of targeted configurations, and even whether a gear ratio of 2 remains the right choice must be re-derived. That re-derivation is the core of this thesis, not a copy of the Delta formulas.",
    )

    add_heading_num(doc, "3. Research gap")
    add_para(doc, "From Nguyen et al. (2020) and related GSM work, three gaps are relevant to this project:", space_after=6)
    add_bullet(
        doc,
        " Existing GSM papers treat serial arms and the spatial Delta. A hanging five-bar needs a new static model and a new notion of targeted configuration (for example, symmetric poses with the end-effector on the midline).",
        bold_lead="The method has not been carried through on a planar closed-loop five-bar.",
    )
    add_bullet(
        doc,
        " The Delta paper already shows that different tasks prefer different target angles, but only compares a few discrete values. There is no systematic, trajectory-based selection rule.",
        bold_lead="The targeted configuration is still chosen by hand.",
    )
    add_bullet(
        doc,
        " A planar five-bar with two GSMs is a realistic prototype. Residual unbalance can be checked mechanically by measuring the force needed to move the mechanism, without electronic sensing.",
        bold_lead="The Delta GSM design was not built.",
    )

    add_heading_num(doc, "4. Aim, objectives, and research questions")
    add_subhead(doc, "4.1 Aim")
    add_para(
        doc,
        "To establish a GSM-based gravity-compensation design for a planar five-bar mechanism, including analytical sizing, numerical evaluation, task-oriented selection of the targeted configuration, and a simple mechanical prototype test.",
    )
    add_subhead(doc, "4.2 Objectives")
    add_bullet(
        doc,
        "Derive the gravitational torques at the two actuated joints of a top-fixed planar five-bar, including link weights and a payload at the end-effector.",
    )
    add_bullet(
        doc,
        "Mount one GSM on each actuated crank and derive the spring-torque model under the same practical assumptions as the Delta paper (free-length initial spring, gear arm much shorter than the connecting rod).",
    )
    add_bullet(
        doc,
        "Redefine targeted configurations for the five-bar and solve for spring stiffness ki and installation angle ψi. Examine whether gear ratio ng = 2 still follows from matching torque forms.",
    )
    add_bullet(
        doc,
        "Evaluate workspace and trajectory performance using torque reduction rate (TRR), mean and peak TRR, gravity compensation density (GCD), and energy reduction rate (ERR) where joint speed is known. These indices remain numerical; they are not measured with sensors on the prototype.",
    )
    add_bullet(
        doc,
        "Given representative planar trajectories, treat the targeted configuration as a design variable and select it by a scan or a small optimization, subject to limits on stiffness, installation angle, and spring stroke. Compare the result with hand-picked targets.",
    )
    add_bullet(
        doc,
        "3D-print a simple five-bar with two GSMs and verify the design with a handheld mechanical force gauge, by comparing the quasi-static force required to move the mechanism with the GSMs engaged and with the GSMs disconnected.",
    )
    add_subhead(doc, "4.3 Research questions")
    add_para(
        doc,
        "RQ1. Can the targeted-configuration approximation used for the Delta be reformulated so that a planar five-bar achieves useful gravity compensation with two GSMs and practical compression springs?",
    )
    add_para(
        doc,
        "RQ2. For a given planar trajectory, does an automatically selected targeted configuration outperform the usual symmetric hand-picked choice in mean torque, peak torque, and energy (numerical indices)?",
    )
    add_para(
        doc,
        "RQ3. On a 3D-printed prototype, does a handheld mechanical force gauge still show a clear drop in the force required to move the mechanism after the GSMs are fitted, despite friction and manufacturing error?",
    )

    add_heading_num(doc, "5. Scope")
    add_subhead(doc, "In scope")
    add_bullet(doc, "Planar five-bar, ground link on top, two swinging actuated cranks, one GSM per crank.")
    add_bullet(doc, "Quasi-static design; dynamics used only afterwards for numerical ERR on specified trajectories.")
    add_bullet(doc, "Task-oriented selection of the targeted configuration (one main design variable if left–right symmetry is kept).")
    add_bullet(
        doc,
        "One simple 3D-printed prototype, validated by handheld mechanical force-gauge tests (unbalanced vs balanced).",
    )
    add_subhead(doc, "Out of scope")
    add_bullet(
        doc,
        "Electronic force or torque sensing, load cells, strain gauges, motor-current estimation, data-acquisition systems, and any other sensing technology.",
    )
    add_bullet(doc, "Reproducing or 3D-printing the spatial Delta in Nguyen et al. (2020).")
    add_bullet(doc, "A full variable-payload adaptive stiffness mechanism (to be discussed only).")
    add_bullet(doc, "Surgical or rehabilitation hardware as the main object.")
    add_bullet(doc, "High-speed dynamic control or industrial-grade metrology.")

    add_heading_num(doc, "6. Experimental validation (mechanical, not electronic)")
    add_para(
        doc,
        "Following the supervisor’s advice, prototype tests will use a simple mechanical principle rather than sensing technology. If a gravity compensator were perfect and frictionless, no external force would be needed to move the mechanism quasi-statically. Any leftover gravitational unbalance, plus friction, appears as the force that must be applied by hand to displace the end-effector. Measuring that force with and without the GSMs is therefore a direct check of the design.",
    )
    add_subhead(doc, "6.1 Instrument")
    add_para(
        doc,
        "The intended instrument is a handheld analog mechanical force gauge of the IMADA PS type, specifically the PS-10N (capacity 10 N, resolution 0.05 N, push/pull, real-time and peak modes, tare ring, no batteries). Product information: https://imada.com/products/ps-10n-mechanical-force-gauge/. Prototype link masses and payload will be chosen so that the forces stay inside this 10 N range. If a trial shows that the unbalanced force exceeds 10 N, a higher-capacity mechanical gauge in the same analog PS family will be used. An electronic gauge, load cell, or motor will not be substituted.",
    )
    add_subhead(doc, "6.2 Procedure")
    add_bullet(doc, "Fix the five-bar in a vertical plane on a rigid frame. Fit a small hook at the end-effector so that the gauge’s hook attachment can be connected in a known direction.")
    add_bullet(doc, "Mark a small set of poses and two slow paths: a vertical move on the midline, and one simple pick-and-place style path in the plane.")
    add_bullet(doc, "Two conditions, same payload: GSMs disconnected or springs unloaded (unbalanced); GSMs engaged (balanced).")
    add_bullet(doc, "Move the end-effector slowly by pulling or pushing through the mechanical gauge, so that inertia is negligible. Use the tare ring after the attachment is fitted. Use peak mode for the largest force along a path, and real-time mode for the force at marked poses.")
    add_bullet(doc, "Repeat each path at least three times. Record the dial reading by hand. No data logger is required.")
    add_bullet(
        doc,
        "Report a force reduction, analogous to the numerical TRR: 1 − F_balanced / F_unbalanced, for peak force and for selected poses. Discuss leftover force caused by friction, backlash, and printed-part mass error.",
    )
    add_para(
        doc,
        "An optional extra check, still purely mechanical, is a hanging-mass test: attach a known small mass at the end-effector and observe whether the balanced mechanism stays at rest at the targeted configuration. This is only a qualitative supplement to the force-gauge measurements.",
    )
    add_para(
        doc,
        "Safety: spring covers, travel limits, no shock loading, and no high-speed tests (mechanical gauges of this type are not intended for destructive or impact loads).",
    )

    add_heading_num(doc, "7. Design flow")
    add_para(
        doc,
        "Five-bar dimensions + GSM geometry + targeted configuration Θ  →  k, ψ  →  workspace indices (TRR, GCD)  →  trajectory indices (M-TRR, P-TRR, ERR)  →  scan Θ for a given task → new k, ψ  →  prototype test with a mechanical force gauge: unbalanced vs balanced.",
    )

    add_heading_num(doc, "8. Proposed thesis structure")
    for item in [
        "Introduction",
        "Literature review (static balancing, Delta/GSM paper, five-bar robots)",
        "Five-bar kinematics and gravitational-torque model",
        "GSM modelling and analytical balancing design",
        "Numerical performance evaluation",
        "Task-oriented selection of the targeted configuration",
        "Prototype fabrication and mechanical force-gauge experiments",
        "Discussion (including variable payload as future work)",
        "Conclusions",
    ]:
        add_bullet(doc, item)
    add_para(
        doc,
        "The task-oriented chapter is part of this thesis. It is not a separate project. Experimental work in the prototype chapter is limited to the mechanical force-gauge procedure in Section 6.",
    )

    add_heading_num(doc, "9. Expected contributions")
    add_bullet(
        doc,
        "A complete GSM gravity-compensation formulation for a top-fixed planar five-bar, including targeted-configuration conditions that are not copied from the Delta.",
    )
    add_bullet(
        doc,
        "Evidence, on this mechanism, that trajectory-based selection of Θ improves numerical torque and energy indices relative to hand-picked targets.",
    )
    add_bullet(
        doc,
        "A simple prototype, with the force required to move the balanced and unbalanced mechanism compared on a handheld mechanical force gauge, and with an honest account of 3D-printing limitations.",
    )
    add_bullet(
        doc,
        "A compact undergraduate thesis that follows the supervisor’s gravity-compensation line without repeating the published Delta case study and without introducing sensing technology.",
    )

    add_heading_num(doc, "10. Feasibility, resources, and risks")
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Item", "Plan"),
        ("Theory / simulation", "MATLAB or Python; 2D kinematics are standard."),
        (
            "Fabrication",
            "UOW 3D printers, bearings, catalogue springs, basic workshop. Link masses sized so that moving forces fit a 10 N mechanical gauge.",
        ),
        (
            "Measurement",
            "Handheld analog mechanical force gauge (IMADA PS-10N or same-family mechanical gauge). No electronic sensors, load cells, or motor-current measurement.",
        ),
        (
            "Risk: printed gear friction hides the effect",
            "Use bearings on the five-bar joints; keep GSM gear force moderate; report leftover force as a measured offset on the gauge.",
        ),
        (
            "Risk: unbalanced force exceeds 10 N",
            "Reduce payload or printed mass first. Only if still needed, switch to a higher-range analog mechanical gauge in the same family, not to an electronic sensor.",
        ),
        (
            "Risk: scope growth",
            "No Delta hardware; no adaptive-stiffness device; no sensing technology; target-angle study stays a 1-D scan unless time remains.",
        ),
    ]
    for i, (a, b) in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], a, bold=(i == 0))
        set_cell_text(table.rows[i].cells[1], b, bold=(i == 0))
    doc.add_paragraph()

    add_heading_num(doc, "11. Timeline")
    add_para(
        doc,
        "The plan assumes one final-year project session (adjust to the official UOW calendar). A mid-project review with the supervisor should freeze geometry, payload, and the set of paths before printing.",
    )
    t2 = doc.add_table(rows=6, cols=3)
    t2.style = "Table Grid"
    trows = [
        ("Phase", "Period", "Work"),
        ("1", "Weeks 1–3", "Literature, five-bar CAD sketch, kinematics and Tw derivation"),
        ("2", "Weeks 4–6", "GSM torque model, targeted-configuration design, baseline simulation"),
        ("3", "Weeks 7–9", "Trajectory scan of Θ, comparison tables and plots"),
        ("4", "Weeks 10–14", "Detail design, print, assemble, mechanical force-gauge tests (PS-10N)"),
        ("5", "Weeks 15–18", "Thesis writing, extra gauge tests, discussion of limitations and variable payload"),
    ]
    for i, (a, b, c) in enumerate(trows):
        set_cell_text(t2.rows[i].cells[0], a, bold=(i == 0))
        set_cell_text(t2.rows[i].cells[1], b, bold=(i == 0))
        set_cell_text(t2.rows[i].cells[2], c, bold=(i == 0))
    doc.add_paragraph()

    add_heading_num(doc, "12. What this proposal is not")
    add_para(
        doc,
        "This is not a reconstruction of the FANUC / theoretical Delta example in Nguyen et al. (2020). The Delta paper is the method template. The new object is the planar five-bar; the new chapter is automatic, task-oriented choice of the targeted configuration; the new evidence is a prototype comparison of the force required to move the mechanism, measured with a handheld mechanical force gauge. Sensing technology is outside the project.",
    )

    add_heading_num(doc, "13. References (starting set)")
    refs = [
        "V. L. Nguyen, C.-Y. Lin and C.-H. Kuo, “Gravity compensation design of Delta parallel robots using gear-spring modules,” Mechanism and Machine Theory, vol. 154, 104046, 2020.",
        "V. L. Nguyen, C.-Y. Lin and C.-H. Kuo, “Gravity compensation design of planar articulated robotic arms using the gear-spring modules,” ASME Journal of Mechanisms and Robotics, vol. 12, no. 3, 031014, 2020.",
        "J. L. Herder, Energy-free Systems: Theory, Conception and Design of Statically Balanced Spring Mechanisms, Ph.D. thesis, TU Delft, 2001.",
        "V. Arakelian, “Gravity compensation in robotics,” Advanced Robotics, vol. 30, no. 2, pp. 79–96, 2016.",
        "I. Simionescu, L. Ciupitu and L. C. Ionita, “Static balancing with elastic systems of DELTA parallel robots,” Mechanism and Machine Theory, vol. 87, pp. 150–162, 2015.",
        "X.-J. Liu, J. Wang and G. Pritschow, “Kinematics, singularity and workspace of planar 5R symmetrical parallel mechanisms,” Mechanism and Machine Theory, vol. 41, no. 2, pp. 145–169, 2006.",
        "IMADA, Inc., “PS-10N Mechanical Force Gauge,” https://imada.com/products/ps-10n-mechanical-force-gauge/ (accessed August 2026).",
    ]
    for i, ref in enumerate(refs, 1):
        add_para(doc, f"[{i}] {ref}", space_after=4)

    add_heading_num(doc, "14. Note to the supervisor")
    add_para(
        doc,
        "This revision follows the advice not to use sensing technology. Prototype validation is now limited to a handheld analog mechanical force gauge (IMADA PS-10N or the same mechanical family) that measures the force required to move the balanced and unbalanced five-bar. Electronic transducers, load cells, and motor-based torque estimation have been removed from the plan.",
    )

    out = Path("/workspace/docs/Project_Proposal_Bo_Zhang_revised.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    artifact = Path("/opt/cursor/artifacts")
    if artifact.is_dir():
        doc.save(artifact / "Project_Proposal_Bo_Zhang_revised.docx")
    print("wrote", out)


if __name__ == "__main__":
    build()
