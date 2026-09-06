#!/usr/bin/env python3
"""Week 1–3: hanging planar five-bar dimensions, IK check, and CAD sketch."""

from __future__ import annotations

import json
import math
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from matplotlib.patches import Arc, Circle, FancyArrowPatch, FancyBboxPatch, Rectangle
from matplotlib.lines import Line2D

# ---------------------------------------------------------------------------
# Baseline geometry (metres unless noted). Symmetric hanging 5R.
# Ground on top, cranks hang down, end-effector at the coupler joint.
# ---------------------------------------------------------------------------
PARAMS = {
    "d": 0.180,  # ground length, joint A to joint B
    "l1": 0.120,  # left crank AP
    "l2": 0.120,  # right crank BQ
    "l3": 0.180,  # left coupler PE
    "l4": 0.180,  # right coupler QE
    "payload_kg": 0.30,
    "m_crank_kg": 0.040,
    "m_coupler_kg": 0.050,
    "s_crank": 0.5,  # COM of crank from actuated joint, fraction of length
    "s_coupler": 0.5,  # COM of coupler from crank–coupler joint
    "shaft_d_mm": 8.0,
    "link_w_mm": 20.0,
    "link_t_mm": 8.0,
    "sample_E": (0.0, -0.180),  # sample pose on the midline
}

G = 9.81
OUT = Path("/workspace/docs/week1-3")
FIG = OUT / "figures"
ART = Path("/opt/cursor/artifacts")


def A_B(d: float):
    return np.array([-d / 2.0, 0.0]), np.array([d / 2.0, 0.0])


def circle_intersections(c1, r1, c2, r2):
    c1 = np.asarray(c1, float)
    c2 = np.asarray(c2, float)
    delta = c2 - c1
    dist = float(np.linalg.norm(delta))
    if dist < 1e-12 or dist > r1 + r2 + 1e-12 or dist < abs(r1 - r2) - 1e-12:
        return []
    a = (r1**2 - r2**2 + dist**2) / (2.0 * dist)
    h2 = r1**2 - a**2
    if h2 < -1e-12:
        return []
    h = math.sqrt(max(h2, 0.0))
    p = c1 + a * delta / dist
    perp = np.array([-delta[1], delta[0]]) / dist
    if h < 1e-12:
        return [p]
    return [p + h * perp, p - h * perp]


def pick_elbow(points, side: str):
    """Prefer the assembly that hangs below the base and opens outward."""
    if not points:
        return None
    pts = sorted(points, key=lambda p: (p[1], p[0] if side == "left" else -p[0]))
    # lowest point first; if tie, more outward
    below = [p for p in pts if p[1] <= 0.02]
    pool = below if below else pts
    if side == "left":
        return min(pool, key=lambda p: p[0])  # more negative x
    return max(pool, key=lambda p: p[0])


def inverse_kinematics(x, y, p=PARAMS):
    A, B = A_B(p["d"])
    E = np.array([x, y], float)
    P_cands = circle_intersections(A, p["l1"], E, p["l3"])
    Q_cands = circle_intersections(B, p["l2"], E, p["l4"])
    P = pick_elbow(P_cands, "left")
    Q = pick_elbow(Q_cands, "right")
    if P is None or Q is None:
        return None
    th_L = math.atan2(P[1] - A[1], P[0] - A[0])
    th_R = math.atan2(Q[1] - B[1], Q[0] - B[0])
    return {
        "A": A,
        "B": B,
        "P": np.asarray(P),
        "Q": np.asarray(Q),
        "E": E,
        "theta_L_rad": th_L,
        "theta_R_rad": th_R,
        "theta_L_deg": math.degrees(th_L),
        "theta_R_deg": math.degrees(th_R),
        "theta_L_from_down_deg": math.degrees(th_L + math.pi / 2),
        "theta_R_from_down_deg": math.degrees(th_R + math.pi / 2),
    }


def reachable_on_midline(p=PARAMS, n=400):
    ys = np.linspace(-0.35, -0.02, n)
    ok = []
    for y in ys:
        if inverse_kinematics(0.0, y, p) is not None:
            ok.append(y)
    return (min(ok), max(ok)) if ok else None


def stretched_singular_y(p=PARAMS):
    """Midline y where crank+coupler are fully stretched (approx. lower reach)."""
    # |AE| = l1+l3, AE from A=(-d/2,0) to (0,y)
    r = p["l1"] + p["l3"]
    half = p["d"] / 2.0
    if r <= half:
        return None
    return -math.sqrt(r**2 - half**2)


def folded_singular_y(p=PARAMS):
    r = abs(p["l3"] - p["l1"])
    half = p["d"] / 2.0
    if r <= half:
        return None
    return -math.sqrt(r**2 - half**2)


def estimate_unbalance_force_N(p=PARAMS):
    """Rough vertical force at EE if all moving mass is lifted (upper bound)."""
    m = 2 * p["m_crank_kg"] + 2 * p["m_coupler_kg"] + p["payload_kg"]
    return m * G, m


def draw_dimensioned_skeleton(sol, p, path: Path, title: str):
    fig, ax = plt.subplots(figsize=(10.5, 8.2), dpi=160)
    A, B, P, Q, E = sol["A"], sol["B"], sol["P"], sol["Q"], sol["E"]

    # frame
    ax.plot([A[0] - 0.04, B[0] + 0.04], [0.02, 0.02], color="#444444", lw=10, solid_capstyle="butt", zorder=1)
    ax.plot([A[0], B[0]], [0, 0], color="#222222", lw=3, zorder=2)

    def bar(p1, p2, color, lw=5):
        ax.plot([p1[0], p2[0]], [p1[1], p2[1]], color=color, lw=lw, solid_capstyle="round", zorder=3)

    bar(A, P, "#1f4e79", 5.5)
    bar(B, Q, "#1f4e79", 5.5)
    bar(P, E, "#c45911", 5.5)
    bar(Q, E, "#c45911", 5.5)

    joints = [("A", A), ("B", B), ("P", P), ("Q", Q), ("E", E)]
    for name, pt in joints:
        ax.plot(pt[0], pt[1], "o", ms=11, color="white", markeredgecolor="#111", markeredgewidth=1.6, zorder=5)
        off = {
            "A": (-0.028, 0.018),
            "B": (0.012, 0.018),
            "P": (-0.032, -0.008),
            "Q": (0.012, -0.008),
            "E": (0.012, -0.022),
        }[name]
        ax.text(pt[0] + off[0], pt[1] + off[1], name, fontsize=12, fontweight="bold")

    # COMs
    C_L = A + p["s_crank"] * (P - A)
    C_R = B + p["s_crank"] * (Q - B)
    C_PL = P + p["s_coupler"] * (E - P)
    C_QR = Q + p["s_coupler"] * (E - Q)
    for c, lab in [(C_L, r"$C_{aL}$"), (C_R, r"$C_{aR}$"), (C_PL, r"$C_{bL}$"), (C_QR, r"$C_{bR}$")]:
        ax.plot(c[0], c[1], "x", ms=8, mew=1.6, color="#548235", zorder=6)
        ax.text(c[0] + 0.006, c[1] + 0.006, lab, fontsize=9, color="#375623")

    # payload hook
    ax.plot([E[0], E[0]], [E[1], E[1] - 0.028], color="#7f6000", lw=1.6)
    ax.plot(E[0], E[1] - 0.028, "v", color="#7f6000", ms=8)
    ax.text(E[0] + 0.012, E[1] - 0.038, r"$m_E$ hook", fontsize=9, color="#7f6000")

    # gravity
    ax.annotate(
        "",
        xy=(0.20, -0.04),
        xytext=(0.20, 0.02),
        arrowprops=dict(arrowstyle="->", color="#c00000", lw=1.6),
    )
    ax.text(0.208, -0.01, "g", color="#c00000", fontsize=12, fontweight="bold")

    # angle arcs from +x
    def angle_arc(origin, ang, label, ysign=-1):
        r = 0.038
        ax.add_patch(Arc(origin, 2 * r, 2 * r, angle=0, theta1=min(0, math.degrees(ang)), theta2=max(0, math.degrees(ang)), color="#7030a0", lw=1.3))
        ax.plot([origin[0], origin[0] + 0.05], [origin[1], origin[1]], color="#888", lw=0.8, ls="--")
        mid = ang / 2.0 if ang < 0 else ang / 2
        ax.text(origin[0] + 0.048 * math.cos(mid), origin[1] + 0.048 * math.sin(mid) + 0.004 * ysign, label, fontsize=10, color="#7030a0")

    angle_arc(A, sol["theta_L_rad"], r"$\theta_L$", 1)
    angle_arc(B, sol["theta_R_rad"], r"$\theta_R$", 1)

    # dimensions
    def dim_h(y, x0, x1, text, color="#111"):
        ax.annotate("", xy=(x1, y), xytext=(x0, y), arrowprops=dict(arrowstyle="<->", color=color, lw=1.1))
        ax.text((x0 + x1) / 2, y + 0.006, text, ha="center", fontsize=10, color=color)

    def dim_v(x, y0, y1, text, color="#111"):
        ax.annotate("", xy=(x, y1), xytext=(x, y0), arrowprops=dict(arrowstyle="<->", color=color, lw=1.1))
        ax.text(x + 0.006, (y0 + y1) / 2, text, va="center", fontsize=10, color=color, rotation=90)

    dim_h(0.045, A[0], B[0], rf"$d = {p['d']*1000:.0f}$ mm")
    # crank length along a parallel offset
    v = P - A
    n = np.array([-v[1], v[0]]) / np.linalg.norm(v)
    off = 0.022 * n
    ax.annotate("", xy=tuple(P + off), xytext=tuple(A + off), arrowprops=dict(arrowstyle="<->", color="#1f4e79", lw=1.1))
    mid = (A + P) / 2 + off * 1.6
    ax.text(mid[0] - 0.03, mid[1], rf"$l_1={p['l1']*1000:.0f}$ mm", fontsize=9, color="#1f4e79")
    v2 = E - P
    n2 = np.array([-v2[1], v2[0]]) / np.linalg.norm(v2)
    off2 = 0.018 * n2
    ax.annotate("", xy=tuple(E + off2), xytext=tuple(P + off2), arrowprops=dict(arrowstyle="<->", color="#c45911", lw=1.1))
    mid2 = (P + E) / 2 + off2 * 1.8
    ax.text(mid2[0] - 0.055, mid2[1] - 0.004, rf"$l_3={p['l3']*1000:.0f}$ mm", fontsize=9, color="#c45911")

    ax.set_aspect("equal")
    ax.set_xlim(-0.28, 0.26)
    ax.set_ylim(-0.30, 0.08)
    ax.set_xlabel("x (m)")
    ax.set_ylabel("y (m)")
    ax.set_title(title, fontsize=13, pad=10)
    ax.grid(True, ls=":", alpha=0.45)
    ax.axhline(0, color="#bbb", lw=0.6)
    ax.axvline(0, color="#bbb", lw=0.6)
    legend = [
        Line2D([0], [0], color="#1f4e79", lw=4, label="cranks $a_L,a_R$ (actuated)"),
        Line2D([0], [0], color="#c45911", lw=4, label="couplers $b_L,b_R$"),
        Line2D([0], [0], color="#548235", marker="x", ls="", label="link COM (mid-link)"),
        Line2D([0], [0], color="#7f6000", marker="v", ls="", label="payload hook at E"),
    ]
    ax.legend(handles=legend, loc="lower left", framealpha=0.95, fontsize=8)
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def draw_layout_schematic(p, path: Path):
    fig, ax = plt.subplots(figsize=(10.2, 6.6), dpi=160)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6.6)
    ax.axis("off")
    ax.set_title("Hanging planar 5R — layout and angle convention (Week 1–3 sketch)", fontsize=13, pad=8)

    # simple cartoon not to scale
    ax.plot([2.0, 8.0], [5.4, 5.4], color="#444", lw=14, solid_capstyle="butt")
    ax.plot([2.6, 7.4], [5.15, 5.15], color="#222", lw=3)
    A = np.array([2.6, 5.15])
    B = np.array([7.4, 5.15])
    P = np.array([3.3, 3.15])
    Q = np.array([6.7, 3.15])
    E = np.array([5.0, 1.55])
    ax.plot([A[0], P[0], E[0], Q[0], B[0]], [A[1], P[1], E[1], Q[1], B[1]], color="#1f4e79", lw=4)
    for pt, name, dx, dy in [
        (A, "A  (actuated)", -0.15, 0.18),
        (B, "B  (actuated)", 0.08, 0.18),
        (P, "P", -0.35, 0.0),
        (Q, "Q", 0.12, 0.0),
        (E, "E  (end-effector)", 0.12, -0.28),
    ]:
        ax.plot(*pt, "o", ms=10, color="white", markeredgecolor="#111", markeredgewidth=1.4)
        ax.text(pt[0] + dx, pt[1] + dy, name, fontsize=10, fontweight="bold")

    ax.annotate("", xy=(5.0, 0.55), xytext=(5.0, 1.15), arrowprops=dict(arrowstyle="->", color="#c00000", lw=1.5))
    ax.text(5.12, 0.7, "g  (down = −y)", color="#c00000", fontsize=10)

    box = (
        "Convention\n"
        "• Origin O: midpoint of ground AB, y = 0 on the ground line\n"
        "• +x to the right, +y upward, gravity −y\n"
        "• θ_L, θ_R: crank angles from +x, counterclockwise\n"
        "• Also report φ = θ + 90°: angle from the downward vertical\n"
        "• Symmetric target poses: x_E = 0 and θ_L + θ_R = −180°\n"
        "  (the two cranks mirror each other about the midline)\n"
        "• GSM not drawn this week — reserved for Week 4"
    )
    ax.text(
        0.35,
        2.55,
        box,
        fontsize=9,
        va="top",
        family="DejaVu Sans",
        bbox=dict(boxstyle="round,pad=0.45", facecolor="#fff8e7", edgecolor="#bfbfbf"),
    )
    ax.text(
        7.55,
        2.55,
        "Do not detail this week\n• gears / springs (GSM)\n• bearings exploded view\n• 3D print fillets\n\nMust freeze before Week 4\n• d, l1…l4\n• payload m_E\n• angle convention",
        fontsize=9,
        va="top",
        bbox=dict(boxstyle="round,pad=0.45", facecolor="#f2f2f2", edgecolor="#bfbfbf"),
    )
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def write_scad(p, path: Path):
    # millimetres in OpenSCAD
    d, l1, l3 = p["d"] * 1000, p["l1"] * 1000, p["l3"] * 1000
    w, t, hole = p["link_w_mm"], p["link_t_mm"], p["shaft_d_mm"] + 0.2
    scad = f"""// Hanging planar five-bar — Week 1–3 kinematic sketch (mm)
// Not a final print file. GSM is omitted.
// Ground on top. Open in OpenSCAD, F5 preview, F6 render.

d = {d:.1f};
l1 = {l1:.1f};
l2 = l1;
l3 = {l3:.1f};
l4 = l3;
w  = {w:.1f};
t  = {t:.1f};
hole = {hole:.1f};

module link(len) {{
    difference() {{
        hull() {{
            translate([0,0,0]) cylinder(h=t, d=w, center=true);
            translate([len,0,0]) cylinder(h=t, d=w, center=true);
        }}
        translate([0,0,0]) cylinder(h=t+2, d=hole, center=true);
        translate([len,0,0]) cylinder(h=t+2, d=hole, center=true);
    }}
}}

module ground() {{
    difference() {{
        hull() {{
            translate([-d/2,0,0]) cylinder(h=t+2, d=w+8, center=true);
            translate([ d/2,0,0]) cylinder(h=t+2, d=w+8, center=true);
        }}
        translate([-d/2,0,0]) cylinder(h=t+6, d=hole, center=true);
        translate([ d/2,0,0]) cylinder(h=t+6, d=hole, center=true);
    }}
}}

// Preview pose: left/right cranks at sample IK (approx.)
theta_L = {math.degrees(math.atan2(-0.180, 0.0)) - 30}; // replace after Week 1–3 IK script
theta_R = 180 - theta_L;

color("dimgray") ground();
translate([-d/2,0,t]) rotate([0,0,theta_L]) color("steelblue") link(l1);
translate([ d/2,0,t]) rotate([0,0,theta_R]) color("steelblue") link(l2);
// Couplers are left as separate parts for the student to assemble after IK is coded.
"""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(scad)


def write_word(p, sol, y_range, f_est, m_tot, fig1, fig2, path: Path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    def font(run, size=12, bold=False):
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)

    def para(doc, text, **kw):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(kw.get("sa", 6))
        if kw.get("center"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        font(r, size=kw.get("size", 12), bold=kw.get("bold", False))
        return p

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    para(doc, "Week 1–3 Deliverable 2", size=12, bold=True, center=True, sa=2)
    para(doc, "Planar Five-Bar CAD Sketch and Parameter Table", size=16, bold=True, center=True, sa=4)
    para(
        doc,
        "Bo Zhang (8571260)  |  Supervisor: Dr Chin-Hsing Kuo  |  September 2026",
        size=11,
        center=True,
        sa=10,
    )

    para(doc, "1. What this week freezes", bold=True, size=14, sa=6)
    para(
        doc,
        "This is a kinematic sketch, not a finished print. Ground link on top, two actuated cranks hanging down, two couplers meeting at the end-effector E. No GSM, gears, or springs are drawn. The numbers below are the baseline that Week 1–3 inverse kinematics and Tw will use. They can still be scaled later, but the ratios and the angle convention should stay.",
    )

    para(doc, "2. Why these numbers", bold=True, size=14, sa=6)
    para(
        doc,
        f"Desktop size (ground d = {p['d']*1000:.0f} mm, cranks {p['l1']*1000:.0f} mm, couplers {p['l3']*1000:.0f} mm) fits a bench frame and leaves space later for one GSM on each crank. The design is left–right symmetric so targeted poses collapse to the midline x = 0. Couplers are longer than cranks, which is the usual 5R choice: a useful workspace below the base without an immediate fully-stretched singularity on the midline.",
    )
    para(
        doc,
        f"The IMADA PS-10N is 10 N. Estimated moving mass is {m_tot:.2f} kg (two cranks, two couplers, payload {p['payload_kg']*1000:.0f} g). If that whole mass were lifted vertically, the force would be about {f_est:.1f} N — inside 10 N, with a little margin for friction. Payload is 300 g, not 1 kg. Printed PLA masses are estimates (20 mm × 8 mm bars plus hubs); they will be replaced by weighed parts before Tw is finalised.",
    )
    if y_range:
        para(
            doc,
            f"On the midline, the outward-elbow inverse kinematics used in the sketch is feasible for y ∈ [{y_range[0]*1000:.0f}, {y_range[1]*1000:.0f}] mm. The sample pose E = (0, −180) mm sits in the middle of that interval and is the first targeted-configuration candidate.",
        )

    para(doc, "3. Parameter table (baseline)", bold=True, size=14, sa=6)
    rows = [
        ("Symbol", "Meaning", "Value", "Note"),
        ("d", "Ground AB", "180 mm", "Top fixed bar; joints A, B"),
        ("l1 = l2", "Crank AP, BQ", "120 mm", "Actuated links"),
        ("l3 = l4", "Coupler PE, QE", "180 mm", "Floating links"),
        ("O", "Origin", "midpoint of AB", "+x right, +y up, g = −y"),
        ("θ_L, θ_R", "Crank angles", "from +x, CCW", "Also store φ = θ + 90° from downward vertical"),
        ("m_a", "Each crank mass", "40 g (est.)", "Replace by weighing"),
        ("m_b", "Each coupler mass", "50 g (est.)", "Replace by weighing"),
        ("s_a, s_b", "COM location", "mid-link", "Week 1–3 assumption"),
        ("m_E", "Payload at E", "300 g", "Hook at E; PS-10N range"),
        ("shaft", "Revolute pins", "8 mm", "Ball bearings later"),
        ("bar section", "Link section", "20 × 8 mm", "3D-print starting size"),
        ("GSM", "Gear-spring module", "not this week", "Mass later added into m_a"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(txt)
            font(r, size=10, bold=(i == 0))
    doc.add_paragraph()

    para(doc, "4. Sample pose used in the drawing", bold=True, size=14, sa=6)
    para(
        doc,
        f"E = (0, −180) mm. Inverse kinematics (outward elbows): "
        f"θ_L = {sol['theta_L_deg']:.1f}°, θ_R = {sol['theta_R_deg']:.1f}° "
        f"(from +x). From the downward vertical: "
        f"φ_L = {sol['theta_L_from_down_deg']:.1f}°, φ_R = {sol['theta_R_from_down_deg']:.1f}°. "
        f"P ≈ ({sol['P'][0]*1000:.1f}, {sol['P'][1]*1000:.1f}) mm, "
        f"Q ≈ ({sol['Q'][0]*1000:.1f}, {sol['Q'][1]*1000:.1f}) mm.",
    )

    para(doc, "5. Drawings", bold=True, size=14, sa=6)
    para(doc, "Figure 1. Dimensioned skeleton at the sample midline pose.", sa=4)
    doc.add_picture(str(fig1), width=Inches(6.3))
    para(doc, "Figure 2. Layout, coordinate frame, and what is frozen this week.", sa=4)
    doc.add_picture(str(fig2), width=Inches(6.3))

    para(doc, "6. Files", bold=True, size=14, sa=6)
    para(
        doc,
        "OpenSCAD sketch (bars and holes only): cad/five_bar_week1_sketch.scad. "
        "Machine-readable parameters: docs/week1-3/five_bar_parameters.json. "
        "Next Week 1–3 task: inverse-kinematics program using these lengths, then Tw.",
    )

    para(doc, "7. What not to change without a reason", bold=True, size=14, sa=6)
    para(
        doc,
        "Keep symmetry (l1 = l2, l3 = l4). Keep payload ≤ 300 g until a trial on the mechanical gauge says otherwise. Do not add GSM mass into Tw until the module size is chosen in Week 4. If a later print needs a scale factor, multiply all of d, l1, l3 by the same factor so the kinematics stay similar.",
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    if ART.is_dir():
        doc.save(ART / path.name)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    FIG.mkdir(parents=True, exist_ok=True)
    p = PARAMS
    sol = inverse_kinematics(*p["sample_E"], p)
    if sol is None:
        raise SystemExit("Sample pose is not reachable — change dimensions.")
    y_range = reachable_on_midline(p)
    f_est, m_tot = estimate_unbalance_force_N(p)

    fig1 = FIG / "five_bar_dimensioned_skeleton.png"
    fig2 = FIG / "five_bar_layout_convention.png"
    draw_dimensioned_skeleton(
        sol,
        p,
        fig1,
        f"Five-bar skeleton at E = (0, −180) mm   |   θ_L = {sol['theta_L_deg']:.1f}°, θ_R = {sol['theta_R_deg']:.1f}°",
    )
    draw_layout_schematic(p, fig2)
    if ART.is_dir():
        import shutil

        shutil.copy(fig1, ART / fig1.name)
        shutil.copy(fig2, ART / fig2.name)

    payload = {
        **{k: v for k, v in p.items() if k != "sample_E"},
        "sample_E_mm": [p["sample_E"][0] * 1000, p["sample_E"][1] * 1000],
        "sample_theta_L_deg": sol["theta_L_deg"],
        "sample_theta_R_deg": sol["theta_R_deg"],
        "midline_y_range_mm": [y_range[0] * 1000, y_range[1] * 1000] if y_range else None,
        "estimated_moving_mass_kg": m_tot,
        "estimated_vertical_unbalance_N": f_est,
        "angle_convention": "theta from +x CCW; phi = theta + 90 deg from downward vertical",
    }
    (OUT / "five_bar_parameters.json").write_text(json.dumps(payload, indent=2))

    write_scad(p, Path("/workspace/cad/five_bar_week1_sketch.scad"))
    write_word(
        p,
        sol,
        y_range,
        f_est,
        m_tot,
        fig1,
        fig2,
        Path("/workspace/docs/FiveBar_CAD_Sketch_Week1-3_Bo_Zhang.docx"),
    )

    print("OK sample pose", {k: sol[k] for k in ("theta_L_deg", "theta_R_deg")})
    print("P", sol["P"], "Q", sol["Q"])
    print("midline y range m", y_range)
    print("est force N", f_est, "mass kg", m_tot)
    print("wrote", fig1)
    print("wrote", fig2)


if __name__ == "__main__":
    main()
